"""Convert combined_tax_report.md to a formatted .docx file."""
import re
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

doc = Document()

# ── Page setup ──
for section in doc.sections:
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

# ── Style defaults ──
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)
style.font.color.rgb = RGBColor(0x2D, 0x2D, 0x2D)
style.paragraph_format.space_after = Pt(6)
style.paragraph_format.line_spacing = 1.15

for level in range(1, 4):
    h = doc.styles[f'Heading {level}']
    h.font.name = 'Calibri'
    h.font.color.rgb = RGBColor(0x00, 0x3D, 0x6B)
    h.font.bold = True
    if level == 1:
        h.font.size = Pt(20)
        h.paragraph_format.space_before = Pt(24)
        h.paragraph_format.space_after = Pt(12)
    elif level == 2:
        h.font.size = Pt(16)
        h.paragraph_format.space_before = Pt(18)
        h.paragraph_format.space_after = Pt(8)
    else:
        h.font.size = Pt(13)
        h.paragraph_format.space_before = Pt(12)
        h.paragraph_format.space_after = Pt(6)

NAVY = RGBColor(0x00, 0x3D, 0x6B)
ACCENT = RGBColor(0x00, 0x6B, 0x77)
DARK = RGBColor(0x2D, 0x2D, 0x2D)
GREY = RGBColor(0x5A, 0x5A, 0x5A)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
HEADER_BG = RGBColor(0x00, 0x3D, 0x6B)
ALT_BG = RGBColor(0xF2, 0xF6, 0xFA)


def set_cell_shading(cell, color):
    """Set background color on a table cell."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    hex_color = str(color)  # RGBColor.__str__ returns hex like '003D6B'
    shading = tcPr.makeelement(qn('w:shd'), {
        qn('w:fill'): hex_color,
        qn('w:val'): 'clear',
    })
    tcPr.append(shading)


def add_styled_table(headers, rows):
    """Add a formatted table to the document."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'

    # Header row
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ''
        p = cell.paragraphs[0]
        run = p.add_run(h)
        run.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = WHITE
        run.font.name = 'Calibri'
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_shading(cell, HEADER_BG)

    # Data rows
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = ''
            p = cell.paragraphs[0]
            run = p.add_run(val.strip())
            run.font.size = Pt(10)
            run.font.name = 'Calibri'
            run.font.color.rgb = DARK
            if r_idx % 2 == 1:
                set_cell_shading(cell, ALT_BG)

    # Column widths — distribute evenly
    total = Inches(6.0)
    col_w = total / len(headers)
    for row in table.rows:
        for cell in row.cells:
            cell.width = int(col_w)

    doc.add_paragraph()  # spacing


def add_run_with_bold(paragraph, text):
    """Parse **bold** markers and add runs accordingly."""
    parts = re.split(r'(\*\*.*?\*\*)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
            run.font.name = 'Calibri'
            run.font.size = Pt(11)
        else:
            run = paragraph.add_run(part)
            run.font.name = 'Calibri'
            run.font.size = Pt(11)


# ── Read the markdown ──
with open('combined_tax_report.md', 'r', encoding='utf-8') as f:
    lines = f.readlines()

i = 0
in_table = False
table_headers = []
table_rows = []

while i < len(lines):
    line = lines[i].rstrip('\n')

    # Skip pure separator lines
    if line.strip() == '---':
        if in_table:
            add_styled_table(table_headers, table_rows)
            in_table = False
            table_headers = []
            table_rows = []
        i += 1
        continue

    # Table detection
    if '|' in line and line.strip().startswith('|'):
        cells = [c.strip() for c in line.strip().strip('|').split('|')]
        # Check if next line is separator
        if not in_table:
            # Check if this is a header row (next line has dashes)
            if i + 1 < len(lines) and re.match(r'^[\s|:-]+$', lines[i + 1].strip()):
                in_table = True
                table_headers = cells
                i += 2  # skip header + separator
                continue
            else:
                # Single row without header - treat as text
                pass
        else:
            # Check if it's a separator row
            if re.match(r'^[\s|:-]+$', line.strip()):
                i += 1
                continue
            table_rows.append(cells)
            # Check if next line is not a table row
            if i + 1 >= len(lines) or '|' not in lines[i + 1]:
                add_styled_table(table_headers, table_rows)
                in_table = False
                table_headers = []
                table_rows = []
            i += 1
            continue

    # Flush any pending table
    if in_table and '|' not in line:
        add_styled_table(table_headers, table_rows)
        in_table = False
        table_headers = []
        table_rows = []

    # Headings
    if line.startswith('# ') and not line.startswith('## '):
        text = line[2:].strip()
        p = doc.add_heading(text, level=1)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        i += 1
        continue
    if line.startswith('## '):
        text = line[3:].strip()
        # Check if next line is also a heading (subtitle)
        if i + 1 < len(lines) and lines[i + 1].startswith('## '):
            p = doc.add_heading(text, level=1)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            i += 1
            # Add subtitle
            subtitle = lines[i].rstrip('\n')[3:].strip()
            p2 = doc.add_paragraph()
            p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p2.add_run(subtitle)
            run.font.size = Pt(13)
            run.font.color.rgb = GREY
            run.font.name = 'Calibri'
            run.italic = True
            i += 1
            continue
        doc.add_heading(text, level=2)
        i += 1
        continue
    if line.startswith('### '):
        doc.add_heading(line[4:].strip(), level=3)
        i += 1
        continue

    # Bullet points
    if line.startswith('- ') or line.startswith('  - ') or line.startswith('    - '):
        indent = 0
        stripped = line
        if line.startswith('    - '):
            indent = 2
            stripped = line[6:]
        elif line.startswith('  - '):
            indent = 1
            stripped = line[4:]
        else:
            stripped = line[2:]

        p = doc.add_paragraph(style='List Bullet')
        p.paragraph_format.left_indent = Inches(0.3 + indent * 0.3)
        add_run_with_bold(p, stripped.strip())
        i += 1
        continue

    # Numbered items
    m = re.match(r'^(\d+)\.\s+(.*)', line)
    if m:
        p = doc.add_paragraph(style='List Number')
        add_run_with_bold(p, m.group(2))
        i += 1
        continue

    # Empty lines
    if not line.strip():
        i += 1
        continue

    # Regular paragraph (handle *italic* and **bold**)
    # Strip leading bold markers for emphasis paragraphs
    p = doc.add_paragraph()
    text = line.strip()

    # Handle italic with *text*
    add_run_with_bold(p, text)

    i += 1

# Flush final table if any
if in_table:
    add_styled_table(table_headers, table_rows)

# ── Save ──
out = 'combined_tax_report.docx'
doc.save(out)
print(f'Saved: {out}')
