"""Key Insights section builder — visual-rich with charts, framework tables, and callouts."""

import logging
import re

from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

from report.styles import (
    add_section_heading, add_subsection_heading, add_sub_subsection_heading,
    add_body_text, add_chart_image, add_formatted_text, clean_markdown,
    add_callout_box, add_kpi_dashboard, add_horizontal_rule,
    add_insight_title, add_slide_divider, add_slide_break,
    add_section_label, add_footer_insight_bar,
    add_chart_with_sidebar, add_cagr_analysis_card, add_numbered_drivers_card,
    add_kpi_cards_with_badge,
    add_visual_bullet_list, add_section_intro_strip,
    set_cell_shading, _set_generous_padding, _keep_table_on_one_page,
    _remove_table_borders,
    NAVY, GOLD, STEEL_GRAY, DARK_TEXT, WHITE,
    DARK_BLUE, ACCENT_BLUE, ACCENT_ORANGE,
    STYLE_CHART_CAPTION, STYLE_BODY,
    FONT_DISPLAY, FONT_BODY, FONT_LABEL,
)
from report.mapper import SectionPlan, get_me_for_dimension, get_all_years, get_cagr_key, get_snapshot_years
from report.tables import add_forecast_table
from report.charts import (
    chart_bubble_attractiveness, chart_horizontal_cagr_ranking, chart_combo_forecast,
    chart_porters_radar, chart_waterfall_growth, chart_market_growth_area,
    chart_impact_heatmap,
)

logger = logging.getLogger(__name__)


# ─── Main Builder ────────────────────────────────────────────────────────────


def build_key_insights(doc: Document, plan: SectionPlan, me_global: dict, content: dict = None):
    """Build the Key Insights section with researched content + visual elements."""
    years = get_all_years(me_global)
    cagr_key = get_cagr_key(me_global)

    # Extract forecast start year from CAGR key (e.g. "cagr_2025_2032" → "2025")
    _ck_parts = cagr_key.split("_") if cagr_key else []
    forecast_start_yr = _ck_parts[1] if len(_ck_parts) >= 2 and _ck_parts[1].isdigit() else None

    # Compute intro metrics — scope info (NOT CAGR, which the metrics dashboard covers)
    sub_count = len(plan.subsections) if plan.subsections else 0
    intro_metrics = []
    if sub_count:
        intro_metrics.append({"value": str(sub_count), "label": "Topics"})
    # Count frameworks (PEST, Porter's, etc.)
    framework_count = 0
    sub_titles = [s.get("title", "").lower() for s in (plan.subsections or [])]
    for t in sub_titles:
        if any(kw in t for kw in ("pest", "porter", "attractiveness", "risk", "pricing", "supply chain")):
            framework_count += 1
    if framework_count:
        intro_metrics.append({"value": str(framework_count), "label": "Frameworks"})
    # Count dynamics children (drivers, restraints, opportunities)
    dynamics_count = 0
    for s in (plan.subsections or []):
        if "dynamic" in s.get("title", "").lower():
            dynamics_count = len(s.get("children", []))
            break
    if dynamics_count:
        intro_metrics.append({"value": str(dynamics_count), "label": "Force Factors"})

    add_section_heading(doc, plan.title or "Key Insights", section_label="KEY INSIGHTS")
    add_section_intro_strip(doc, "Key Insights",
                            "Market dynamics, competitive forces, and strategic frameworks",
                            icon="◎", metrics=intro_metrics or None)

    subsections_content = content.get("subsections", {}) if content else {}

    # Key Metrics Dashboard at top of section
    _add_key_metrics_box(doc, me_global, years, cagr_key)

    # Executive summary callout (if content has it)
    exec_summary = content.get("executive_summary", "") if content else ""
    total = me_global.get("total", {})
    val_forecast = total.get("value", {}).get("forecast", {})
    val_cagr = total.get("value", {}).get(cagr_key, 0)
    if not exec_summary:
        if val_forecast and years:
            try:
                base_val = float(val_forecast.get(years[0], 0))
                end_val = float(val_forecast.get(years[-1], 0))
                cagr_pct = float(val_cagr) * 100
                exec_summary = (
                    f"The market is projected to grow from **US$ {base_val:,.1f} Mn** "
                    f"in {years[0]} to **US$ {end_val:,.1f} Mn** by {years[-1]}, "
                    f"representing a CAGR of **{cagr_pct:.1f}%** during the forecast period. "
                    f"This section provides detailed analysis of market dynamics, "
                    f"competitive forces, and strategic frameworks driving this growth trajectory."
                )
            except (ValueError, TypeError):
                pass
    if exec_summary:
        add_callout_box(doc, exec_summary, box_type="key_finding", title="Executive Summary")
        # Footer insight bar summarizing market trajectory
        try:
            if val_forecast and years:
                base_v = float(val_forecast.get(years[0], 0))
                end_v = float(val_forecast.get(years[-1], 0))
                growth_pct = ((end_v / base_v) - 1) * 100 if base_v > 0 else 0
                add_footer_insight_bar(doc,
                    f"The market is expected to grow {growth_pct:.0f}% over the forecast period, "
                    f"driven by strong fundamentals across key segments and geographies.")
        except (ValueError, TypeError):
            pass

    # Market Growth Trajectory — area chart with CAGR + drivers sidebar
    if val_forecast and years:
        add_slide_break(doc)  # New slide for growth trajectory
        try:
            base_f = float(val_forecast.get(years[0], 0))
            end_f = float(val_forecast.get(years[-1], 0))
            cagr_f = float(val_cagr) * 100
            cagr_period = cagr_key.replace("cagr_", "").replace("_", "-")

            # Area chart — the hero visual (sidebar-optimized size)
            img = chart_market_growth_area(
                f"Market Growth Trajectory ({cagr_period})",
                val_forecast, years, cagr_pct=cagr_f,
                fig_size=(6.0, 3.8), forecast_start_yr=forecast_start_yr,
            )

            # Sidebar: CAGR card + top drivers list
            val_unit = me_global.get("market_value", {}).get("unit", "Mn")
            start_label = f"US$ {base_f:,.1f} {val_unit} ({years[0]})"
            end_label = f"US$ {end_f:,.1f} {val_unit} ({years[-1]})"
            drivers_raw = content.get("drivers", []) if content else []
            if not drivers_raw and content:
                # Try to extract from subsections content
                for sub_title, sub_c in content.get("subsections", {}).items():
                    if "dynamic" in sub_title.lower() or "driver" in sub_title.lower():
                        if isinstance(sub_c, dict) and isinstance(sub_c.get("structured"), list):
                            for item in sub_c["structured"]:
                                if item.get("type", "").lower() == "driver":
                                    drivers_raw.append(item.get("factor", ""))
                        break

            def build_growth_sidebar(cell):
                add_cagr_analysis_card(cell, f"{cagr_f:.1f}%", start_label, end_label,
                                       label=f"CAGR ({cagr_period})")
                if drivers_raw:
                    add_numbered_drivers_card(cell, [
                        {"title": str(d)[:70], "desc": ""}
                        for d in drivers_raw[:3]
                    ], title="KEY GROWTH DRIVERS")

            add_chart_with_sidebar(doc, img, build_growth_sidebar,
                                   chart_width=5.8, sidebar_width=3.7,
                                   caption="Total addressable market value progression with CAGR")

            # Waterfall chart — market value breakdown by segment
            add_slide_break(doc)  # New slide for waterfall
            val_section = me_global.get("market_value", {})
            seg_contribs = {}
            for dim_key, items in val_section.items():
                if dim_key.startswith("by_") and isinstance(items, dict) and dim_key != "by_region":
                    for name, data in items.items():
                        try:
                            contrib = float(data.get("forecast", {}).get(years[-1], 0)) - \
                                      float(data.get("forecast", {}).get(years[0], 0))
                            seg_contribs[name] = contrib
                        except (ValueError, TypeError):
                            continue
                    if seg_contribs:
                        break

            img = chart_waterfall_growth(
                f"Market Value Progression: {years[0]} to {years[-1]}",
                base_f, end_f, years[0], years[-1],
                segments=seg_contribs if seg_contribs else None,
                unit=val_unit,
            )
            add_chart_image(doc, img, caption="Waterfall breakdown of market growth drivers")

        except Exception as e:
            logger.warning(f"Growth trajectory charts failed: {e}")

    # Iterate subsections
    for sub in plan.subsections:
        title = sub.get("title", "")
        # New slide for each subsection
        add_slide_break(doc)

        # Add section label based on subsection type
        title_lower = title.lower()
        if "dynamic" in title_lower or "impact" in title_lower:
            add_section_label(doc, "MARKET DYNAMICS")
        elif "pest" in title_lower:
            add_section_label(doc, "PEST ANALYSIS")
        elif "porter" in title_lower:
            add_section_label(doc, "COMPETITIVE FORCES")
        elif "attractiveness" in title_lower:
            add_section_label(doc, "MARKET ATTRACTIVENESS")
        elif "supply" in title_lower:
            add_section_label(doc, "SUPPLY CHAIN")
        elif "development" in title_lower:
            add_section_label(doc, "KEY DEVELOPMENTS")
        elif "pricing" in title_lower:
            add_section_label(doc, "PRICING ANALYSIS")
        add_subsection_heading(doc, title)

        # Extract content and structured data
        sub_content = subsections_content.get(title, {})
        text_content = ""
        structured = {}
        if isinstance(sub_content, dict):
            text_content = sub_content.get("text", "")
            structured = sub_content.get("structured", {})
        elif isinstance(sub_content, str) and sub_content:
            text_content = sub_content

        # Render children headings for ALL subsections so TOC children appear in the report.
        # Skip only when the text content already mentions the child title (dedup).
        children = sub.get("children", [])
        for child in children:
            child_title = child.get("title", "")
            if not child_title:
                continue
            # Skip if the LLM-generated text already includes this child heading
            if text_content and child_title.lower() in text_content.lower():
                continue
            add_sub_subsection_heading(doc, child_title)

        # Render narrative text as visual bullets (strip structured blocks first)
        if text_content:
            clean_text = _strip_structured_blocks(text_content)
            # Strip markdown tables when structured data will generate styled visuals
            if structured:
                clean_text = _strip_markdown_tables(clean_text)
            add_formatted_text(doc, clean_text)

        # ── Visual elements based on subsection type ──────────────────
        title_lower = title.lower()

        # Market Dynamics → Impact Heatmap + Analysis table
        if "dynamic" in title_lower and isinstance(structured, list) and structured:
            # Filter out "challenge" type — only driver, opportunity, restraint
            structured = [item for item in structured
                          if item.get("type", "").lower().strip() != "challenge"]

            try:
                img = chart_impact_heatmap("Strategic Factor Analysis", structured)
                add_chart_image(doc, img, caption="Market dynamics positioned by type and impact severity")
            except Exception as e:
                logger.warning(f"Impact heatmap chart failed: {e}")

            _add_impact_analysis_table(doc, structured)
            # Insight callout summarizing key drivers
            drivers = [clean_markdown(d.get("factor", "")) for d in structured
                       if d.get("type", "").lower() == "driver"]
            if drivers:
                driver_list = ", ".join(drivers[:3])
                add_callout_box(doc,
                    f"**Key Growth Drivers:** {driver_list}. "
                    f"These factors are expected to significantly influence market trajectory "
                    f"during the forecast period.",
                    box_type="key_finding", title="Market Drivers Summary")

        # PEST Analysis → 2x2 colored grid
        if "pest" in title_lower and isinstance(structured, dict) and structured:
            _add_pest_grid(doc, structured)

        # Porter's Five Forces → radar chart + styled rating table
        if "porter" in title_lower and isinstance(structured, dict) and structured:
            # Radar chart first (visual impact)
            try:
                img = chart_porters_radar("Porter's Five Forces Analysis", structured)
                add_chart_image(doc, img, caption="Porter's Five Forces — Competitive Intensity Map")
            except Exception as e:
                logger.warning(f"Porter's radar chart failed: {e}")
            # Then the detailed table
            _add_porters_table(doc, structured)
            # Callout with overall assessment
            high_forces = [clean_markdown(f) for f, d in structured.items()
                          if d.get("rating", "").lower().strip() == "high"]
            if high_forces:
                forces_str = ", ".join(high_forces)
                add_callout_box(doc,
                    f"**High competitive pressure** identified in: {forces_str}. "
                    f"These forces require strategic attention and may compress margins "
                    f"for market participants.",
                    box_type="insight", title="Competitive Alert")

        # Market Attractiveness → bubble chart + CAGR ranking
        if "attractiveness" in title_lower:
            _add_attractiveness_visuals(doc, me_global, years, cagr_key)

        # Pricing Analysis → chart + tables from ME data
        if "pricing" in title_lower:
            _add_pricing_visuals(doc, me_global, years, cagr_key)

        # Supply Chain → flow diagram table
        if "supply" in title_lower and "chain" in title_lower:
            stages = structured if isinstance(structured, list) and structured else None
            _add_supply_chain_flow(doc, stages)

        # Key Developments → timeline table
        if "development" in title_lower and isinstance(structured, list) and structured:
            _add_developments_timeline(doc, structured)

        # Market Attractiveness fallback (no content)
        if "attractiveness" in title_lower and not sub_content:
            add_body_text(doc,
                "Market attractiveness analysis evaluates segments based on market size, "
                "growth rate, and competitive intensity to identify high-potential investment areas."
            )

        # (Slide divider removed — the next subsection's add_slide_break handles
        #  page separation, and the divider was causing extra empty pages.)


# ─── Key Metrics Callout Box ─────────────────────────────────────────────────


def _add_key_metrics_box(doc: Document, me_global: dict, years: list[str], cagr_key: str):
    """Add a KPI dashboard showing insight-oriented metrics distinct from the Overview KPIs."""
    if not years:
        return

    first_yr = years[0]
    last_yr = years[-1]
    total = me_global.get("total", {})
    val_forecast = total.get("value", {}).get("forecast", {})
    val_cagr = total.get("value", {}).get(cagr_key, 0)

    try:
        base_f = float(val_forecast.get(first_yr, 0))
        end_f = float(val_forecast.get(last_yr, 0))
        cagr_f = float(val_cagr) * 100
    except (ValueError, TypeError):
        return

    cagr_period = cagr_key.replace("cagr_", "").replace("_", "-")

    # 1. Market addition — absolute growth over the period
    addition = end_f - base_f
    if addition >= 1000:
        addition_str = f"US$ {addition / 1000:,.2f} Bn"
    else:
        addition_str = f"US$ {addition:,.1f} Mn"

    # 2. Growth multiple (e.g. 2.1× expansion)
    multiple = end_f / base_f if base_f > 0 else 0
    multiple_str = f"{multiple:.1f}×"

    # 3. Fastest growing segment (highest CAGR across any by_* dimension)
    fastest_seg, fastest_cagr = _find_fastest_growing(me_global.get("market_volume", {}), cagr_key)

    # 4. Leading region
    leading_region = _find_largest_item(me_global.get("market_volume", {}), last_yr, only_key="by_region")

    metrics = [
        {
            "label": f"Market Addition ({first_yr}–{last_yr})",
            "value": addition_str,
            "subtitle": "Absolute Value Growth",
        },
        {
            "label": "Growth Multiple",
            "value": multiple_str,
            "subtitle": f"Market expansion by {last_yr}",
            "badge": f"▲ {cagr_f:.1f}% CAGR",
            "badge_color": "#00BCD4",
        },
        {
            "label": "Fastest Growing Segment",
            "value": fastest_seg[:18] if fastest_seg else "—",
            "subtitle": f"{fastest_cagr:.1f}% CAGR" if fastest_cagr else "Top growth driver",
        },
        {
            "label": "Leading Region",
            "value": leading_region[:18] if leading_region else "—",
            "subtitle": "By market volume share",
        },
    ]

    add_kpi_cards_with_badge(doc, metrics)


def _find_fastest_growing(section: dict, cagr_key: str) -> tuple:
    """Find the segment with the highest CAGR across all by_* dimensions.

    Returns (segment_name, cagr_as_float_percent) or ("", 0) if not found.
    """
    best_name, best_cagr = "", 0.0
    for dim_key, items in section.items():
        if not dim_key.startswith("by_") or not isinstance(items, dict):
            continue
        for name, data in items.items():
            try:
                cagr_val = float(data.get(cagr_key, 0)) * 100
            except (ValueError, TypeError):
                continue
            if cagr_val > best_cagr:
                best_cagr = cagr_val
                best_name = name
    return best_name, best_cagr


def _find_largest_item(section: dict, year: str, exclude_keys: tuple = (), only_key: str = None) -> str:
    """Find the item with highest forecast value in a given year across dimensions."""
    best_name, best_val = "", 0

    for dim_key, items in section.items():
        if dim_key in exclude_keys or not isinstance(items, dict):
            continue
        if only_key and dim_key != only_key:
            continue

        for name, data in items.items():
            try:
                val = float(data.get("forecast", {}).get(year, 0))
            except (ValueError, TypeError):
                continue
            if val > best_val:
                best_val = val
                best_name = name

    return best_name


# ─── Impact Analysis Table ───────────────────────────────────────────────────


def _add_colored_left_border(cell, color_hex: str):
    """Add a thick colored left border to a table cell."""
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    tcPr = cell._element.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), "24")  # 3pt
    left.set(qn("w:space"), "0")
    left.set(qn("w:color"), color_hex)
    tcBorders.append(left)
    tcPr.append(tcBorders)


def _add_impact_analysis_table(doc: Document, items: list[dict]):
    """Render impact analysis as a color-coded table matching the project navy/gold theme."""
    if not items:
        return

    from report.tables import _set_row_cant_split, _set_row_keep_with_next, _apply_light_borders
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn as _qn

    caption = doc.add_paragraph("Impact Analysis Summary", style=STYLE_CHART_CAPTION)
    caption.paragraph_format.keep_with_next = True

    # Type config — teal + green palette (3 types only: driver, opportunity, restraint)
    type_config = {
        "driver":      {"icon": "\u25B2", "border": "006B77", "label_color": NAVY},
        "restraint":   {"icon": "\u25BC", "border": "2E7D32", "label_color": RGBColor(0x2E, 0x7D, 0x32)},
        "opportunity": {"icon": "\u2605", "border": "009688", "label_color": RGBColor(0x00, 0x96, 0x88)},
    }
    impact_config = {
        "high":     {"icon": "\u2B24\u2B24\u2B24", "bg": "006B77", "text": WHITE},   # Dark Teal
        "medium":   {"icon": "\u2B24\u2B24\u25CB",  "bg": "009688", "text": WHITE},  # Teal-Green
        "moderate": {"icon": "\u2B24\u2B24\u25CB",  "bg": "009688", "text": WHITE},  # Teal-Green
        "low":      {"icon": "\u2B24\u25CB\u25CB",  "bg": "2E7D32", "text": WHITE},  # Forest Green
    }
    # Alternating row shading — matches project tables
    ROW_EVEN = "E0F0F0"   # blue-tint
    ROW_ODD  = "FFFFFF"   # white

    n_rows = 1 + len(items)
    table = doc.add_table(rows=n_rows, cols=4)
    table.style = "Table Grid"
    _keep_table_on_one_page(table)
    table.alignment = 1
    _apply_light_borders(table)

    # Header row — navy, keep with first data row
    _set_row_cant_split(table.rows[0])
    _set_row_keep_with_next(table.rows[0])
    headers = ["Factor", "Type", "Impact Level", "Time Horizon"]
    for j, h in enumerate(headers):
        cell = table.cell(0, j)
        cell.text = h
        set_cell_shading(cell, "006B77")
        _set_cell_padding(cell, 6)
        for para in cell.paragraphs:
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in para.runs:
                run.font.color.rgb = WHITE
                run.font.bold = True
                run.font.size = Pt(9)
                run.font.name = FONT_BODY

    for i, item in enumerate(items):
        row_idx = i + 1
        factor = clean_markdown(item.get("factor", ""))
        type_val = clean_markdown(item.get("type", ""))
        impact = clean_markdown(item.get("impact", ""))
        horizon = clean_markdown(item.get("horizon", ""))

        type_key = type_val.lower().strip()
        tc = type_config.get(type_key, type_config["driver"])
        row_bg = ROW_EVEN if i % 2 == 0 else ROW_ODD

        _set_row_cant_split(table.rows[row_idx])

        # Factor cell — alternating row bg + colored left accent border
        cell = table.cell(row_idx, 0)
        cell.text = factor
        _add_colored_left_border(cell, tc["border"])
        set_cell_shading(cell, row_bg)
        _set_cell_padding(cell, 6)
        for para in cell.paragraphs:
            for run in para.runs:
                run.font.size = Pt(9)
                run.font.name = FONT_BODY
                run.font.bold = True
                run.font.color.rgb = DARK_TEXT

        # Type cell — icon + label, colored text on row bg
        cell = table.cell(row_idx, 1)
        cell.text = ""
        set_cell_shading(cell, row_bg)
        _set_cell_padding(cell, 6)
        para = cell.paragraphs[0]
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        icon_run = para.add_run(f"{tc['icon']} ")
        icon_run.font.size = Pt(10)
        icon_run.font.color.rgb = tc["label_color"]
        icon_run.font.name = FONT_BODY
        label_run = para.add_run(type_val.upper())
        label_run.font.size = Pt(8)
        label_run.font.bold = True
        label_run.font.color.rgb = tc["label_color"]
        label_run.font.name = FONT_BODY

        # Impact cell — bold colored badge (saturated, not pastel)
        cell = table.cell(row_idx, 2)
        cell.text = ""
        _set_cell_padding(cell, 6)
        impact_key = impact.lower().strip()
        ic = impact_config.get(impact_key, {"icon": "\u2B24", "bg": "5D6D7E", "text": WHITE})
        set_cell_shading(cell, ic["bg"])
        para = cell.paragraphs[0]
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        dots_run = para.add_run(f"{ic['icon']}  ")
        dots_run.font.size = Pt(6)
        dots_run.font.color.rgb = ic["text"]
        dots_run.font.name = FONT_BODY
        label_run = para.add_run(impact.upper())
        label_run.font.size = Pt(9)
        label_run.font.bold = True
        label_run.font.color.rgb = ic["text"]
        label_run.font.name = FONT_BODY

        # Horizon cell — standard alternating bg
        cell = table.cell(row_idx, 3)
        cell.text = horizon
        set_cell_shading(cell, row_bg)
        _set_cell_padding(cell, 6)
        for para in cell.paragraphs:
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in para.runs:
                run.font.size = Pt(9)
                run.font.name = FONT_BODY
                run.font.color.rgb = DARK_TEXT

    doc.add_paragraph("")


# ─── PEST Analysis Grid ──────────────────────────────────────────────────────


def _set_cell_padding(cell, pt_val: int = 8):
    """Add padding to a table cell using XML (w:tcMar)."""
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    tcPr = cell._element.get_or_add_tcPr()
    tcMar = OxmlElement("w:tcMar")
    for edge in ("top", "bottom", "left", "right"):
        elem = OxmlElement(f"w:{edge}")
        elem.set(qn("w:w"), str(pt_val * 20))  # twips (1pt = 20 twips)
        elem.set(qn("w:type"), "dxa")
        tcMar.append(elem)
    tcPr.append(tcMar)


def _add_pest_grid(doc: Document, pest_data: dict):
    """Render PEST analysis as a dramatic 2x2 colored grid with icons, bullets, and large headers."""
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    categories = [
        ("political", "POLITICAL", "006B77", "\U0001F3DB"),        # Dark Teal
        ("economic", "ECONOMIC", "009688", "\U0001F4B0"),           # Teal-Green
        ("social", "SOCIAL", "2E7D32", "\U0001F465"),               # Forest Green
        ("technological", "TECHNOLOGICAL", "00BCD4", "\U0001F4A1"), # Bright Cyan
    ]

    # Fallback icons if emoji not supported
    FALLBACK_ICONS = {
        "political": "\u2696",     # ⚖
        "economic": "\u2B24",      # ⬤
        "social": "\u263A",        # ☺
        "technological": "\u2699", # ⚙
    }

    if not any(pest_data.get(c[0]) for c in categories):
        return

    doc.add_paragraph("PEST Analysis Framework", style=STYLE_CHART_CAPTION)

    table = doc.add_table(rows=2, cols=2)
    table.style = "Table Grid"
    _keep_table_on_one_page(table)
    table.alignment = 1

    # Remove all default borders and set subtle outer border
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else OxmlElement("w:tblPr")
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        elem = OxmlElement(f"w:{edge}")
        elem.set(qn("w:val"), "single")
        elem.set(qn("w:sz"), "8" if edge in ("insideH", "insideV") else "4")
        elem.set(qn("w:color"), "FFFFFF")
        elem.set(qn("w:space"), "0")
        borders.append(elem)
    tblPr.append(borders)

    for idx, (key, label, color, icon) in enumerate(categories):
        row = idx // 2
        col = idx % 2
        cell = table.cell(row, col)

        summary = clean_markdown(pest_data.get(key, "Analysis pending."))
        cell.text = ""
        set_cell_shading(cell, color)
        _set_cell_padding(cell, 12)

        # Header paragraph with icon
        para = cell.paragraphs[0]
        para.paragraph_format.space_after = Pt(8)

        # Use fallback icon (safer for docx)
        fallback = FALLBACK_ICONS.get(key, "\u2022")
        icon_run = para.add_run(f"{fallback}  ")
        icon_run.font.size = Pt(16)
        icon_run.font.color.rgb = WHITE
        icon_run.font.name = "Calibri"

        header_run = para.add_run(label)
        header_run.font.bold = True
        header_run.font.size = Pt(14)
        header_run.font.color.rgb = WHITE
        header_run.font.name = "Calibri"

        # Thin white separator line
        sep_para = cell.add_paragraph()
        sep_para.paragraph_format.space_before = Pt(2)
        sep_para.paragraph_format.space_after = Pt(6)
        sep_run = sep_para.add_run("\u2500" * 40)
        sep_run.font.size = Pt(6)
        sep_run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        sep_run.font.name = "Calibri"

        # Content — split into bullet points if possible
        sentences = [s.strip() for s in summary.replace(". ", ".\n").split("\n") if s.strip()]
        if len(sentences) > 1:
            for sentence in sentences[:5]:  # max 5 bullets
                bp = cell.add_paragraph()
                bp.paragraph_format.space_before = Pt(1)
                bp.paragraph_format.space_after = Pt(2)
                bp.paragraph_format.left_indent = Cm(0.3)

                bullet_run = bp.add_run("\u25B8  ")  # ▸
                bullet_run.font.size = Pt(9)
                bullet_run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                bullet_run.font.name = "Calibri"

                text_run = bp.add_run(sentence.rstrip(".") + ".")
                text_run.font.size = Pt(9)
                text_run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                text_run.font.name = "Calibri"
        else:
            body_para = cell.add_paragraph()
            body_para.paragraph_format.space_before = Pt(2)
            text_run = body_para.add_run(summary)
            text_run.font.size = Pt(10)
            text_run.font.color.rgb = WHITE
            text_run.font.name = "Calibri"

    doc.add_paragraph("")


# ─── Porter's Five Forces Table ──────────────────────────────────────────────


def _add_porters_table(doc: Document, porters_data: dict):
    """Render Porter's Five Forces as a table — project navy/gold/steel theme."""
    from report.tables import _set_row_cant_split, _set_row_keep_with_next, _apply_light_borders

    if not porters_data:
        return

    caption = doc.add_paragraph("Porter's Five Forces — Detailed Assessment", style=STYLE_CHART_CAPTION)
    caption.paragraph_format.keep_with_next = True

    # All colors from project palette only
    rating_config = {
        "high":     {"badge_bg": "006B77", "border": "006B77", "score": 5, "label": "HIGH THREAT",
                     "bar_color": RGBColor(0x00, 0x6B, 0x77)},
        "moderate": {"badge_bg": "009688", "border": "009688", "score": 3, "label": "MODERATE",
                     "bar_color": RGBColor(0x00, 0x96, 0x88)},
        "medium":   {"badge_bg": "009688", "border": "009688", "score": 3, "label": "MODERATE",
                     "bar_color": RGBColor(0x00, 0x96, 0x88)},
        "low":      {"badge_bg": "2E7D32", "border": "2E7D32", "score": 1, "label": "LOW THREAT",
                     "bar_color": RGBColor(0x2E, 0x7D, 0x32)},
    }
    ROW_EVEN = "E0F0F0"
    ROW_ODD  = "FFFFFF"

    forces = list(porters_data.items())
    n_rows = 1 + len(forces)
    table = doc.add_table(rows=n_rows, cols=4)
    table.style = "Table Grid"
    _keep_table_on_one_page(table)
    table.alignment = 1
    _apply_light_borders(table)

    # Header row
    _set_row_cant_split(table.rows[0])
    _set_row_keep_with_next(table.rows[0])
    for j, h in enumerate(["Competitive Force", "Threat Level", "Key Factor", "Intensity"]):
        cell = table.cell(0, j)
        cell.text = h
        set_cell_shading(cell, "006B77")
        _set_cell_padding(cell, 8)
        for para in cell.paragraphs:
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in para.runs:
                run.font.color.rgb = WHITE
                run.font.bold = True
                run.font.size = Pt(10)
                run.font.name = FONT_BODY

    for i, (force, data) in enumerate(forces):
        row_idx = i + 1
        rating = clean_markdown(data.get("rating", "")).lower().strip()
        key_factor = clean_markdown(data.get("key_factor", ""))
        rc = rating_config.get(rating, rating_config["moderate"])
        row_bg = ROW_EVEN if i % 2 == 0 else ROW_ODD

        _set_row_cant_split(table.rows[row_idx])

        # Force name — alternating row bg + colored left accent border
        cell = table.cell(row_idx, 0)
        cell.text = ""
        set_cell_shading(cell, row_bg)
        _set_cell_padding(cell, 8)
        _add_colored_left_border(cell, rc["border"])
        para = cell.paragraphs[0]
        name_run = para.add_run(clean_markdown(force))
        name_run.font.bold = True
        name_run.font.size = Pt(10)
        name_run.font.name = FONT_BODY
        name_run.font.color.rgb = DARK_TEXT

        # Rating — solid badge in theme color
        cell = table.cell(row_idx, 1)
        cell.text = ""
        set_cell_shading(cell, rc["badge_bg"])
        _set_cell_padding(cell, 8)
        para = cell.paragraphs[0]
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        label_run = para.add_run(rc["label"])
        label_run.font.bold = True
        label_run.font.size = Pt(9)
        label_run.font.color.rgb = WHITE
        label_run.font.name = FONT_BODY

        # Key factor
        cell = table.cell(row_idx, 2)
        cell.text = key_factor
        set_cell_shading(cell, row_bg)
        _set_cell_padding(cell, 8)
        for para in cell.paragraphs:
            for run in para.runs:
                run.font.size = Pt(9)
                run.font.name = FONT_BODY
                run.font.color.rgb = DARK_TEXT

        # Intensity — score bar using theme colors
        cell = table.cell(row_idx, 3)
        cell.text = ""
        set_cell_shading(cell, row_bg)
        _set_cell_padding(cell, 8)
        para = cell.paragraphs[0]
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        score = rc["score"]
        filled = score
        empty = 5 - filled
        if filled > 0:
            r = para.add_run("\u2588" * filled)
            r.font.color.rgb = rc["bar_color"]
            r.font.size = Pt(14)
            r.font.name = FONT_BODY
        if empty > 0:
            r = para.add_run("\u2591" * empty)
            r.font.color.rgb = RGBColor(0xD9, 0xD9, 0xD9)
            r.font.size = Pt(14)
            r.font.name = FONT_BODY

    doc.add_paragraph("")


# ─── Market Attractiveness Visuals ───────────────────────────────────────────


def _add_attractiveness_visuals(doc: Document, me_global: dict, years: list[str], cagr_key: str):
    """Add Market Attractiveness bubble chart and CAGR ranking chart."""
    if not years:
        return

    # Find the dimension with the most items for the bubble chart
    vol_section = me_global.get("market_volume", {})
    best_items = {}
    for key, items in vol_section.items():
        if key.startswith("by_") and isinstance(items, dict) and len(items) > len(best_items):
            best_items = items

    if best_items:
        try:
            img = chart_bubble_attractiveness(
                "Market Attractiveness — Segment Analysis",
                best_items, cagr_key,
                base_year=years[0], end_year=years[-1],
            )
            add_chart_image(doc, img, caption="Bubble size represents incremental growth opportunity")
        except Exception as e:
            logger.warning(f"Bubble chart failed: {e}")

    # CAGR ranking across ALL dimensions
    all_cagr = {}
    for key, items in vol_section.items():
        if key.startswith("by_") and isinstance(items, dict):
            for name, data in items.items():
                cagr = data.get(cagr_key, 0)
                if cagr:
                    try:
                        all_cagr[name] = {cagr_key: float(cagr)}
                    except (ValueError, TypeError):
                        continue

    if all_cagr:
        try:
            img = chart_horizontal_cagr_ranking(
                "Segment CAGR Ranking (2025-2032)",
                all_cagr, cagr_key, top_n=12,
            )
            add_chart_image(doc, img, caption="Fastest Growing Segments by CAGR")
        except Exception as e:
            logger.warning(f"CAGR ranking chart failed: {e}")


# ─── Pricing Visuals ─────────────────────────────────────────────────────────


def _add_pricing_visuals(doc: Document, me_global: dict, years: list[str], cagr_key: str):
    """Add pricing forecast table and chart from ME data."""
    pricing_data = me_global.get("pricing", {})
    unit = pricing_data.get("unit", "")

    for dim_key, items in pricing_data.items():
        if dim_key == "unit" or not isinstance(items, dict):
            continue
        dim_name = dim_key.replace("by_", "").replace("_", " ").title()

        # Pricing chart
        try:
            img = chart_combo_forecast(
                f"Pricing Forecast by {dim_name}",
                items, years, value_label="Price", bar_unit=unit, show_yoy_line=False,
            )
            add_chart_image(doc, img, caption=f"Pricing Forecast by {dim_name}")
        except Exception as e:
            logger.warning(f"Pricing chart failed: {e}")

        # Pricing table
        add_forecast_table(doc, f"Pricing by {dim_name}", items, years, cagr_key, unit=unit)


# ─── Supply Chain Flow ───────────────────────────────────────────────────────


_STAGE_DESCRIPTIONS = {
    "raw materials": "Sourcing of base chemicals, polymers, and raw inputs",
    "supplier": "Component and material supply chain partners",
    "manufacturer": "Production, assembly, and quality control",
    "distributor": "Logistics, warehousing, and channel management",
    "end user": "Final consumers, institutions, and applications",
    "retailer": "Point-of-sale, online platforms, and outlets",
    "research": "R&D, formulation, and innovation pipeline",
    "packaging": "Packaging design, labeling, and compliance",
}


def _get_stage_description(stage_name: str) -> str:
    """Get a brief description for a supply chain stage."""
    name_lower = stage_name.lower()
    for key, desc in _STAGE_DESCRIPTIONS.items():
        if key in name_lower:
            return desc
    return "Key value chain participant"


def _add_supply_chain_flow(doc: Document, stages: list[str] = None):
    """Render supply chain as a dramatic flow diagram with numbered stages and gradient colors."""
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    if not stages:
        stages = ["Raw Materials Suppliers", "Manufacturers", "Distributors", "End Users"]

    doc.add_paragraph("Supply Chain Value Flow", style=STYLE_CHART_CAPTION)

    # PPTX-grade gradient: teal / green / cyan alternation
    gradient_colors = [
        "006B77", "009688", "00BCD4", "2E7D32", "5A7D8C", "4DB6AC", "006B77", "009688"
    ]

    # Table: stage cells interleaved with arrow cells, 3 rows (number, name, description)
    n_cols = len(stages) * 2 - 1
    table = doc.add_table(rows=3, cols=n_cols)
    table.style = "Table Grid"
    _keep_table_on_one_page(table)
    table.alignment = 1

    # Remove default borders — use white gap borders for clean separation
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else OxmlElement("w:tblPr")
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        elem = OxmlElement(f"w:{edge}")
        elem.set(qn("w:val"), "single")
        elem.set(qn("w:sz"), "8")
        elem.set(qn("w:color"), "FFFFFF")
        elem.set(qn("w:space"), "0")
        borders.append(elem)
    tblPr.append(borders)

    for i in range(n_cols):
        if i % 2 == 0:
            stage_idx = i // 2
            color = gradient_colors[stage_idx % len(gradient_colors)]
            stage_name = clean_markdown(stages[stage_idx])
            desc = _get_stage_description(stages[stage_idx])

            # Row 0: Step number (circle badge)
            cell = table.cell(0, i)
            cell.text = ""
            set_cell_shading(cell, color)
            _set_cell_padding(cell, 6)
            para = cell.paragraphs[0]
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            num_run = para.add_run(f"STEP {stage_idx + 1}")
            num_run.font.size = Pt(8)
            num_run.font.bold = True
            num_run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            num_run.font.name = "Calibri"

            # Row 1: Stage name (large, bold, white on colored bg)
            cell = table.cell(1, i)
            cell.text = ""
            set_cell_shading(cell, color)
            _set_cell_padding(cell, 10)
            para = cell.paragraphs[0]
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            name_run = para.add_run(stage_name)
            name_run.font.bold = True
            name_run.font.size = Pt(11)
            name_run.font.color.rgb = WHITE
            name_run.font.name = "Calibri"

            # Row 2: Description (light bg, italic)
            cell = table.cell(2, i)
            cell.text = ""
            set_cell_shading(cell, "D8EDED")
            _set_cell_padding(cell, 8)
            para = cell.paragraphs[0]
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            desc_run = para.add_run(desc)
            desc_run.font.size = Pt(8)
            desc_run.font.italic = True
            desc_run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
            desc_run.font.name = "Calibri"
        else:
            # Arrow columns — large flowing arrows
            for row in range(3):
                cell = table.cell(row, i)
                cell.text = ""
                if row == 1:
                    # Big arrow in the middle row
                    set_cell_shading(cell, "FFFFFF")
                    para = cell.paragraphs[0]
                    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    arrow_run = para.add_run("\u27A4")  # ➤
                    arrow_run.font.size = Pt(20)
                    arrow_run.font.bold = True
                    arrow_run.font.color.rgb = GOLD
                    arrow_run.font.name = "Calibri"
                else:
                    set_cell_shading(cell, "FFFFFF")

    doc.add_paragraph("")


# ─── Key Developments Timeline ───────────────────────────────────────────────


def _add_timeline_spine_border(cell):
    """Add a thick left blue border to create a timeline spine effect."""
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    tcPr = cell._element.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), "24")       # thick (3pt)
    left.set(qn("w:space"), "0")
    left.set(qn("w:color"), "00BCD4")  # gold accent
    tcBorders.append(left)
    tcPr.append(tcBorders)


def _add_developments_timeline(doc: Document, items: list[dict]):
    """Render key developments as a dramatic timeline with date badges and card-style rows."""
    if not items:
        return

    doc.add_paragraph("Key Developments Timeline", style=STYLE_CHART_CAPTION)

    n_rows = 1 + len(items)
    table = doc.add_table(rows=n_rows, cols=3)
    table.style = "Table Grid"
    _keep_table_on_one_page(table)
    table.alignment = 1

    # Header — navy PPTX style
    for j, h in enumerate(["Date", "Company", "Development"]):
        cell = table.cell(0, j)
        cell.text = h
        set_cell_shading(cell, "006B77")
        _set_cell_padding(cell, 8)
        for para in cell.paragraphs:
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in para.runs:
                run.font.color.rgb = WHITE
                run.font.bold = True
                run.font.size = Pt(10)
                run.font.name = FONT_BODY

    # PPTX-grade alternating teal/green/cyan date badges
    date_colors = ["006B77", "009688", "2E7D32", "00BCD4"]

    for i, item in enumerate(items):
        row_idx = i + 1
        date_color = date_colors[i % len(date_colors)]

        # Date cell — colored badge style
        cell = table.cell(row_idx, 0)
        cell.text = ""
        set_cell_shading(cell, date_color)
        _set_cell_padding(cell, 8)
        _add_timeline_spine_border(cell)
        para = cell.paragraphs[0]
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # White dot marker
        dot_run = para.add_run("\u25CF ")
        dot_run.font.color.rgb = WHITE
        dot_run.font.size = Pt(10)
        dot_run.font.name = "Calibri"

        # Date text in white
        date_run = para.add_run(clean_markdown(item.get("date", "")))
        date_run.font.bold = True
        date_run.font.size = Pt(10)
        date_run.font.name = "Calibri"
        date_run.font.color.rgb = WHITE

        # Company — bold, with accent bar
        cell = table.cell(row_idx, 1)
        cell.text = ""
        bg_color = "F0F7F7" if row_idx % 2 == 0 else "FFFFFF"
        set_cell_shading(cell, bg_color)
        _set_cell_padding(cell, 8)
        para = cell.paragraphs[0]
        company_run = para.add_run(clean_markdown(item.get("company", "")))
        company_run.font.bold = True
        company_run.font.size = Pt(10)
        company_run.font.name = "Calibri"
        company_run.font.color.rgb = DARK_BLUE

        # Description — with subtle background
        cell = table.cell(row_idx, 2)
        cell.text = ""
        set_cell_shading(cell, bg_color)
        _set_cell_padding(cell, 8)
        para = cell.paragraphs[0]
        desc_run = para.add_run(clean_markdown(item.get("description", "")))
        desc_run.font.size = Pt(9)
        desc_run.font.name = "Calibri"
        desc_run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)

    doc.add_paragraph("")


# ─── Text Rendering Helpers ──────────────────────────────────────────────────


def _strip_structured_blocks(text: str) -> str:
    """Remove ===BLOCK=== structured data sections from LLM text."""
    return re.sub(r'===\w+===(.*?)===END_\w+===', '', text, flags=re.DOTALL).strip()


def _strip_markdown_tables(text: str) -> str:
    """Remove markdown table blocks AND their preceding headings from text.

    Detects contiguous blocks of lines starting with '|' (including separator rows)
    and removes them. Also removes the heading immediately above a table block
    (e.g. "### Summary Table") to avoid orphaned headings on empty pages.
    """
    paragraphs = text.split("\n\n")
    filtered = []
    for i, para in enumerate(paragraphs):
        para_stripped = para.strip()
        # Skip if this paragraph block is a markdown table
        if para_stripped and para_stripped.startswith("|"):
            lines = [l.strip() for l in para_stripped.split("\n") if l.strip()]
            if all(l.startswith("|") for l in lines):
                # Also remove the preceding paragraph if it's a heading or short label
                if filtered:
                    prev = filtered[-1].strip()
                    if (prev.startswith("#") or
                        prev.startswith("**") and prev.endswith("**") or
                        (len(prev) < 60 and not prev.endswith("."))):
                        filtered.pop()
                continue  # Skip entire table block
        filtered.append(para)
    return "\n\n".join(filtered)


