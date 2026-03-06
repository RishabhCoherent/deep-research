"""Appendix / Intro section builder — Research Objectives, Analyst, Methodology."""

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

from report.styles import (
    add_section_heading, add_subsection_heading, add_sub_subsection_heading,
    add_body_text, add_formatted_text,
    add_footer_insight_bar, add_kpi_card_row,
    add_slide_break,
    add_visual_bullet_list, add_section_intro_strip,
    set_cell_shading, _remove_table_borders, _set_generous_padding,
    _keep_table_on_one_page,
    NAVY, GOLD, STEEL_GRAY, OFF_WHITE, DARK_TEXT, WHITE,
    FONT_DISPLAY, FONT_BODY, FONT_LABEL,
)
from report.mapper import SectionPlan


_SUBSECTION_ICONS = {
    "objective": "◎",
    "assumption": "✦",
    "abbreviation": "≡",
    "methodology": "⬡",
    "about": "◈",
    "analyst": "◉",
    "recommendation": "▲",
}


def _get_icon(title: str) -> str:
    title_lower = title.lower()
    for key, icon in _SUBSECTION_ICONS.items():
        if key in title_lower:
            return icon
    return "▸"


def _add_intro_card_row(doc: Document, subsections: list):
    """Add a visual card row showing subsection overview."""
    if not subsections:
        return

    n = min(len(subsections), 4)
    items = subsections[:n]

    table = doc.add_table(rows=1, cols=n)
    _remove_table_borders(table)
    _keep_table_on_one_page(table)
    table.alignment = 1

    bg_colors = ["006B77", "E8F4F5", "009688", "E8F4F5"]
    title_colors_dark = [GOLD, NAVY, GOLD, NAVY]
    text_colors_dark = [OFF_WHITE, DARK_TEXT, OFF_WHITE, DARK_TEXT]

    for j, sub in enumerate(items):
        cell = table.cell(0, j)
        cell.text = ""
        is_dark = j % 2 == 0
        bg = bg_colors[j % len(bg_colors)]
        set_cell_shading(cell, bg)
        _set_generous_padding(cell, top=200, bottom=200, left=160, right=160)

        sub_title = sub.get("title", "")
        icon = _get_icon(sub_title)
        t_color = WHITE if is_dark else NAVY

        para = cell.paragraphs[0]
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        icon_run = para.add_run(f"{icon}")
        icon_run.font.name = FONT_BODY
        icon_run.font.size = Pt(22)
        icon_run.font.color.rgb = t_color

        title_para = cell.add_paragraph()
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_para.paragraph_format.space_before = Pt(8)
        title_para.paragraph_format.space_after = Pt(4)
        title_run = title_para.add_run(sub_title)
        title_run.font.name = FONT_DISPLAY
        title_run.font.size = Pt(11)
        title_run.font.bold = True
        title_run.font.color.rgb = t_color

    doc.add_paragraph("")


def build_appendix_section(doc: Document, plan: SectionPlan, content: dict = None):
    """Build an appendix-type section with premium card layout for intro sections."""
    title = plan.title or f"Section {plan.section_number}"
    subsections = plan.subsections
    sub_titles_lower = [s.get("title", "").lower() for s in subsections]

    if any("methodology" in t for t in sub_titles_lower):
        label = "RESEARCH METHODOLOGY"
    elif any("objective" in t or "abbreviation" in t for t in sub_titles_lower):
        label = "INTRODUCTION"
    elif any("recommendation" in t or "analyst" in t for t in sub_titles_lower):
        label = "ANALYST RECOMMENDATIONS"
    else:
        label = "APPENDIX"

    intro_metrics = []
    if subsections:
        intro_metrics.append({"value": str(len(subsections)), "label": "Topics"})

    add_section_heading(doc, title, section_label=label)
    add_section_intro_strip(doc, title, "Research framework, methodology, and key assumptions",
                            icon="◈", metrics=intro_metrics or None)

    if subsections:
        _add_intro_card_row(doc, subsections)

    subsections_content = content.get("subsections", {}) if content else {}

    for sub in subsections:
        sub_title = sub.get("title", "")
        add_slide_break(doc)  # New slide for each subsection
        add_subsection_heading(doc, sub_title)

        children = sub.get("children", [])
        sub_text = subsections_content.get(sub_title, "")

        if sub_text:
            for child in children:
                child_title = child.get("title", "")
                if child_title:
                    add_sub_subsection_heading(doc, child_title)
            add_formatted_text(doc, sub_text)
        else:
            for child in children:
                add_sub_subsection_heading(doc, child.get("title", ""))

            title_lower = sub_title.lower()
            if "methodology" in title_lower:
                _add_methodology_visual(doc)
            elif "objective" in title_lower:
                _add_objectives_visual(doc)
            elif "abbreviation" in title_lower:
                _add_abbreviations_placeholder(doc)
            elif "analyst" in title_lower or "recommendation" in title_lower:
                add_visual_bullet_list(doc, [
                    "**Investment Opportunities** — Identify high-growth segments and regions",
                    "**Strategic Recommendations** — Actions for existing players and new entrants",
                    "**Market Watch** — Key factors to monitor over the forecast period"
                ], title="Analyst Recommendations")


def _add_methodology_visual(doc: Document):
    """Add a visual 4-step methodology overview card row."""
    steps = [
        {"title": "Secondary Research",
         "desc": "Annual reports, industry databases, and published studies"},
        {"title": "Primary Research",
         "desc": "Expert interviews, surveys, and market participant validation"},
        {"title": "Data Triangulation",
         "desc": "Cross-validation of multiple sources for accuracy"},
        {"title": "Forecasting",
         "desc": "Statistical modelling with CAGR-based projections to 2032"},
    ]

    n = len(steps)
    table = doc.add_table(rows=1, cols=n)
    _remove_table_borders(table)
    _keep_table_on_one_page(table)
    table.alignment = 1
    colors = ["006B77", "009688", "2E7D32", "00BCD4"]

    for j, step in enumerate(steps):
        cell = table.cell(0, j)
        cell.text = ""
        set_cell_shading(cell, colors[j])
        _set_generous_padding(cell, top=160, bottom=160, left=120, right=120)

        para = cell.paragraphs[0]
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        num_run = para.add_run(f"STEP {j + 1}")
        num_run.font.name = FONT_LABEL
        num_run.font.size = Pt(8)
        num_run.font.bold = True
        num_run.font.color.rgb = WHITE

        title_para = cell.add_paragraph()
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_para.paragraph_format.space_before = Pt(6)
        title_para.paragraph_format.space_after = Pt(4)
        title_run = title_para.add_run(step["title"])
        title_run.font.name = FONT_DISPLAY
        title_run.font.size = Pt(11)
        title_run.font.bold = True
        title_run.font.color.rgb = WHITE

        desc_para = cell.add_paragraph()
        desc_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        desc_run = desc_para.add_run(step["desc"])
        desc_run.font.name = FONT_BODY
        desc_run.font.size = Pt(9)
        desc_run.font.color.rgb = WHITE

    doc.add_paragraph("")
    add_footer_insight_bar(doc,
        "This research combines rigorous primary and secondary research methodologies, "
        "validated through expert consultations and industry participant interviews.")


def _add_objectives_visual(doc: Document):
    """Add branded numbered research objectives table."""
    objectives = [
        "Estimate and forecast market size by value and volume across all key segments and geographies",
        "Identify key market drivers, restraints, and opportunities shaping market dynamics",
        "Analyze competitive landscape and profile leading market participants globally",
        "Provide strategic insights for market entry, expansion, and investment decisions",
    ]

    table = doc.add_table(rows=len(objectives), cols=2)
    _remove_table_borders(table)
    _keep_table_on_one_page(table)

    for i, obj in enumerate(objectives):
        num_cell = table.cell(i, 0)
        num_cell.text = ""
        set_cell_shading(num_cell, "006B77" if i % 2 == 0 else "009688")
        _set_generous_padding(num_cell, top=120, bottom=120, left=120, right=80)
        para = num_cell.paragraphs[0]
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = para.add_run(f"0{i + 1}")
        run.font.name = FONT_DISPLAY
        run.font.size = Pt(18)
        run.font.bold = True
        run.font.color.rgb = WHITE

        obj_cell = table.cell(i, 1)
        obj_cell.text = ""
        set_cell_shading(obj_cell, "F0F7F7")
        _set_generous_padding(obj_cell, top=120, bottom=120, left=140, right=100)
        para = obj_cell.paragraphs[0]
        run = para.add_run(obj)
        run.font.name = FONT_BODY
        run.font.size = Pt(10)
        run.font.color.rgb = DARK_TEXT

    doc.add_paragraph("")


def _add_abbreviations_placeholder(doc: Document):
    """Add a styled abbreviations table placeholder."""
    abbrevs = [
        ("CAGR", "Compound Annual Growth Rate"),
        ("ME", "Market Estimation"),
        ("PEST", "Political, Economic, Social, Technological"),
        ("RFID", "Radio-Frequency Identification"),
        ("GPR", "Ground-Penetrating Radar"),
        ("TOC", "Table of Contents"),
    ]

    table = doc.add_table(rows=len(abbrevs) + 1, cols=2)
    table.style = "Table Grid"
    _keep_table_on_one_page(table)
    table.alignment = 1

    for j, h in enumerate(["Abbreviation", "Full Form"]):
        cell = table.cell(0, j)
        cell.text = h
        set_cell_shading(cell, "006B77")
        para = cell.paragraphs[0]
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in para.runs:
            run.font.color.rgb = WHITE
            run.font.bold = True
            run.font.size = Pt(10)
            run.font.name = FONT_BODY

    for i, (abbr, full) in enumerate(abbrevs):
        for j, text in enumerate([abbr, full]):
            cell = table.cell(i + 1, j)
            cell.text = text
            if i % 2 == 0:
                set_cell_shading(cell, "F0F7F7")
            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.size = Pt(9)
                    run.font.name = FONT_BODY
                    if j == 0:
                        run.font.bold = True
                        run.font.color.rgb = NAVY

    doc.add_paragraph("")
