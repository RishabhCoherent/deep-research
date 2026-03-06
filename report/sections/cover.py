"""Cover page + PPTX-style Table of Contents with section cards."""

import re
from datetime import datetime

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from report.styles import (
    STYLE_REPORT_TITLE, STYLE_COVER_SUBTITLE, STYLE_BODY,
    NAVY, GOLD, STEEL_GRAY, OFF_WHITE, DARK_TEXT, WHITE, LIGHT_TEAL_GREEN,
    TEAL_GREEN,
    add_horizontal_rule, set_cell_shading, _remove_table_borders, _set_generous_padding,
    _keep_table_on_one_page,
    _gradient_colors, _set_table_full_bleed, _set_table_grid, _set_cell_width_xml,
    _build_slide_header, _build_slide_footer,
    FONT_DISPLAY, FONT_BODY, FONT_LABEL,
)
from report.mapper import SectionPlan


# ─── Analysis type dots — maps section_type to which analysis methods apply ───
# Each tuple: (secondary, primary, proprietary)  True/False
_ANALYSIS_TYPES = {
    "overview":     (True, True, True),
    "key_insights": (True, True, True),
    "segment":      (True, True, True),
    "region":       (True, True, True),
    "competitive":  (True, True, True),
    "appendix":     (True, False, True),
}

_DOT_COLORS = [
    ("006B77", "Secondary"),    # Navy — secondary research
    ("009688", "Primary"),      # Teal — primary research
    ("00BCD4", "Proprietary"),  # Cyan — proprietary analysis
]


def build_cover(doc: Document, report_title: str, subtitle: str, plans: list[SectionPlan],
                me_data: dict = None):
    """Build the cover page followed by a professional Table of Contents."""
    # Cover page — header shows "Coherent Market Insights" on left
    first_section = doc.sections[0]
    first_section.top_margin = Inches(1.35)
    first_section.bottom_margin = Inches(0.6)
    first_section.header_distance = Inches(0.0)
    first_section.footer_distance = Inches(0.0)
    _build_slide_header(first_section, "Coherent Market Insights", report_title)
    _build_slide_footer(first_section, report_title)

    _build_cover_page(doc, report_title, subtitle, me_data=me_data)

    # TOC — separate section so header shows "Table of Contents (TOC)"
    _build_toc_pages(doc, plans, report_title)


# ─── Cover Page ──────────────────────────────────────────────────────────────


def _build_cover_page(doc: Document, report_title: str, subtitle: str, me_data: dict = None):
    """Professional cover page — header gradient bar is provided by the section header."""
    page_width_dxa = 15840
    margin_dxa = 1080

    # Spacer (header gradient already provides the top bar)
    for _ in range(3):
        doc.add_paragraph("")

    # ── Gold ALL-CAPS label above title ───────────────────────────────
    label_para = doc.add_paragraph()
    label_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    label_para.paragraph_format.space_after = Pt(6)
    label_run = label_para.add_run("MARKET RESEARCH REPORT")
    label_run.font.name = FONT_LABEL
    label_run.font.size = Pt(11)
    label_run.font.bold = True
    label_run.font.color.rgb = GOLD

    # Letter spacing
    rPr = label_run._element.get_or_add_rPr()
    spacing = OxmlElement("w:spacing")
    spacing.set(qn("w:val"), "60")
    rPr.append(spacing)

    # ── Report title ──────────────────────────────────────────────────
    doc.add_paragraph(report_title, style=STYLE_REPORT_TITLE)

    # ── Gold accent line under title ──────────────────────────────────
    add_horizontal_rule(doc, color=GOLD, thickness=6)

    # ── Subtitle with full year range from ME data ────────────────────
    if subtitle and me_data:
        # Replace any partial year range (e.g. 2025-2032) with full range (e.g. 2020-2032)
        global_data = me_data.get("global", {})
        total = global_data.get("total", {})
        val_fc = total.get("value", {}).get("forecast", {})
        all_years = sorted([y for y in val_fc if y.isdigit()], key=int)
        if all_years:
            full_range = f"{all_years[0]}-{all_years[-1]}"
            # Replace any existing year range in subtitle
            subtitle = re.sub(r'\d{4}\s*[-–]\s*\d{4}', full_range, subtitle)
        doc.add_paragraph(subtitle, style=STYLE_COVER_SUBTITLE)
    elif subtitle:
        doc.add_paragraph(subtitle, style=STYLE_COVER_SUBTITLE)

    # ── Date line ─────────────────────────────────────────────────────
    date_para = doc.add_paragraph()
    date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    date_run = date_para.add_run(datetime.now().strftime("%B %Y"))
    date_run.font.size = Pt(12)
    date_run.font.color.rgb = STEEL_GRAY
    date_run.font.name = FONT_LABEL


def _add_cover_stats_strip(doc: Document, me_data: dict):
    """Add a 4-card KPI strip to the cover page showing key market metrics."""
    global_data = me_data.get("global", {})
    total = global_data.get("total", {})
    val_data = total.get("value", {})

    # Find CAGR key
    cagr_key = ""
    for k in val_data.keys():
        if k.startswith("cagr_"):
            cagr_key = k
            break

    val_fc = val_data.get("forecast", {})
    years = sorted([y for y in val_fc if y.isdigit()], key=int)

    if not years:
        for _ in range(3):
            doc.add_paragraph("")
        return

    first_yr, last_yr = years[0], years[-1]

    try:
        base_val = float(val_fc.get(first_yr, 0))
        end_val = float(val_fc.get(last_yr, 0))
        cagr_raw = float(val_data.get(cagr_key, 0)) * 100

        def _fmt(v):
            if v >= 1000:
                return f"US$ {v / 1000:,.2f} Bn"
            return f"US$ {v:,.1f} Mn"

        base_str = _fmt(base_val)
        end_str = _fmt(end_val)
        cagr_str = f"{cagr_raw:.1f}%"

        # Count segments from vol section
        vol_section = global_data.get("market_volume", {})
        n_segments = 0
        for k, v in vol_section.items():
            if k.startswith("by_") and isinstance(v, dict):
                n_segments = max(n_segments, len(v))
        seg_str = f"{n_segments}+" if n_segments else "—"

    except (ValueError, TypeError):
        for _ in range(3):
            doc.add_paragraph("")
        return

    doc.add_paragraph("")

    # 4-card strip — navy background cards with gold values
    table = doc.add_table(rows=1, cols=4)
    _remove_table_borders(table)
    _keep_table_on_one_page(table)
    table.alignment = 1  # CENTER

    stats = [
        {"label": f"Market Value ({first_yr})", "value": base_str, "icon": "◉"},
        {"label": f"Projected Value ({last_yr})", "value": end_str, "icon": "▲"},
        {"label": f"CAGR ({cagr_key.replace('cagr_', '').replace('_', '-')})",
         "value": cagr_str, "icon": "↗"},
        {"label": "Key Segments", "value": seg_str, "icon": "◈"},
    ]

    _cover_card_colors = ["006B77", "009688", "006B77", "009688"]
    for j, stat in enumerate(stats):
        cell = table.cell(0, j)
        cell.text = ""
        set_cell_shading(cell, _cover_card_colors[j % len(_cover_card_colors)])
        _set_generous_padding(cell, top=160, bottom=160, left=120, right=120)

        # Icon + label
        para = cell.paragraphs[0]
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        icon_run = para.add_run(f"{stat['icon']}  ")
        icon_run.font.name = FONT_BODY
        icon_run.font.size = Pt(10)
        icon_run.font.color.rgb = LIGHT_TEAL_GREEN

        label_run = para.add_run(stat["label"].upper())
        label_run.font.name = FONT_LABEL
        label_run.font.size = Pt(7)
        label_run.font.bold = True
        label_run.font.color.rgb = WHITE

        # Value
        val_para = cell.add_paragraph()
        val_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        val_para.paragraph_format.space_before = Pt(4)
        val_run = val_para.add_run(stat["value"])
        val_run.font.name = FONT_DISPLAY
        val_run.font.size = Pt(16)
        val_run.font.bold = True
        val_run.font.color.rgb = WHITE

    doc.add_paragraph("")


# ─── PPTX-Style Table of Contents ────────────────────────────────────────────


def _build_toc_pages(doc: Document, plans: list[SectionPlan], report_title: str = ""):
    """Build a PPTX-style Table of Contents with section cards and analysis-type badges."""
    from docx.enum.section import WD_SECTION
    # Create a new Word section so the header shows "Table of Contents (TOC)"
    toc_section = doc.add_section(WD_SECTION.NEW_PAGE)
    toc_section.orientation = doc.sections[0].orientation
    toc_section.page_width = Inches(11)
    toc_section.page_height = Inches(8.5)
    toc_section.left_margin = Inches(0.75)
    toc_section.right_margin = Inches(0.75)
    toc_section.top_margin = Inches(1.35)
    toc_section.bottom_margin = Inches(0.6)
    toc_section.header_distance = Inches(0.0)
    toc_section.footer_distance = Inches(0.0)
    _build_slide_header(toc_section, "Table of Contents (TOC)", report_title)
    _build_slide_footer(toc_section, report_title)

    # ── TOC Header strip ─────────────────────────────────────────────
    toc_header = doc.add_table(rows=1, cols=2)
    _remove_table_borders(toc_header)
    toc_header.alignment = 1
    _set_table_grid(toc_header, [11600, 4240])

    # Left: title
    left = toc_header.cell(0, 0)
    left.text = ""
    set_cell_shading(left, "006B77")
    _set_generous_padding(left, top=200, bottom=200, left=300, right=100)
    para = left.paragraphs[0]
    run = para.add_run("TABLE OF CONTENTS")
    run.font.name = FONT_DISPLAY
    run.font.size = Pt(20)
    run.font.bold = True
    run.font.color.rgb = WHITE
    # Letter spacing
    rPr = run._element.get_or_add_rPr()
    sp = OxmlElement("w:spacing")
    sp.set(qn("w:val"), "60")
    rPr.append(sp)

    # Right: analysis type legend
    right = toc_header.cell(0, 1)
    right.text = ""
    set_cell_shading(right, "004D56")
    _set_generous_padding(right, top=200, bottom=200, left=160, right=200)
    right.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    legend_para = right.paragraphs[0]
    legend_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for i, (hex_c, label) in enumerate(_DOT_COLORS):
        if i > 0:
            sep = legend_para.add_run("    ")
            sep.font.size = Pt(8)
        dot = legend_para.add_run("●")
        dot.font.size = Pt(10)
        dot.font.color.rgb = RGBColor(int(hex_c[:2], 16), int(hex_c[2:4], 16), int(hex_c[4:], 16))
        lbl = legend_para.add_run(f" {label}")
        lbl.font.name = FONT_LABEL
        lbl.font.size = Pt(8)
        lbl.font.color.rgb = OFF_WHITE

    doc.add_paragraph("")  # small spacer

    # ── Section cards ────────────────────────────────────────────────
    sorted_plans = sorted(plans, key=lambda p: p.section_number)

    for plan_idx, plan in enumerate(sorted_plans):
        # Page break after every ~5 sections to avoid overflow
        if plan_idx > 0 and plan_idx % 5 == 0:
            doc.add_page_break()

        sec_num = plan.section_number
        sec_title = _clean_section_title(plan.title or f"Section {sec_num}")
        analysis = _ANALYSIS_TYPES.get(plan.section_type, (True, False, True))

        # Collect all entries for this section
        entries = _collect_toc_entries(plan)

        # Build a card table: 2 rows × 2 cols
        #   Row 0: section header (navy) | analysis dots
        #   Row 1: entries content      | (merged with row 0 right cell)
        n_entries = len(entries)
        n_rows = 1 + max(1, n_entries)  # header row + entry rows

        table = doc.add_table(rows=n_rows, cols=2)
        _remove_table_borders(table)
        _keep_table_on_one_page(table)
        table.alignment = 1
        _set_table_grid(table, [12600, 3240])

        # ── Header row: "Section N" ──
        hdr_left = table.cell(0, 0)
        hdr_left.text = ""
        set_cell_shading(hdr_left, "006B77")
        _set_generous_padding(hdr_left, top=100, bottom=100, left=200, right=100)
        hp = hdr_left.paragraphs[0]
        hr = hp.add_run(f"Section {sec_num}")
        hr.font.name = FONT_DISPLAY
        hr.font.size = Pt(11)
        hr.font.bold = True
        hr.font.color.rgb = WHITE

        # Header right: analysis dots
        hdr_right = table.cell(0, 1)
        hdr_right.text = ""
        set_cell_shading(hdr_right, "006B77")
        _set_generous_padding(hdr_right, top=100, bottom=100, left=80, right=120)
        hdr_right.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        dp = hdr_right.paragraphs[0]
        dp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for d_idx, (hex_c, _label) in enumerate(_DOT_COLORS):
            if d_idx > 0:
                sp_run = dp.add_run("  ")
                sp_run.font.size = Pt(10)
            dot_run = dp.add_run("●")
            dot_run.font.size = Pt(12)
            if analysis[d_idx]:
                dot_run.font.color.rgb = RGBColor(
                    int(hex_c[:2], 16), int(hex_c[2:4], 16), int(hex_c[4:], 16))
            else:
                dot_run.font.color.rgb = RGBColor(0x33, 0x44, 0x4D)  # dim/inactive

        # ── Entry rows ──
        # First entry row: section title (bold, slightly larger)
        title_cell = table.cell(1, 0)
        title_cell.text = ""
        set_cell_shading(title_cell, "F0F7F7")
        _set_generous_padding(title_cell, top=80, bottom=40, left=240, right=100)
        tp = title_cell.paragraphs[0]
        tr = tp.add_run(sec_title)
        tr.font.name = FONT_DISPLAY
        tr.font.size = Pt(10)
        tr.font.bold = True
        tr.font.color.rgb = NAVY

        # Right cell for first entry row — light bg
        right_cell_1 = table.cell(1, 1)
        right_cell_1.text = ""
        set_cell_shading(right_cell_1, "F0F7F7")

        # Subsequent entry rows: subsections and children
        for e_idx, entry in enumerate(entries):
            row_idx = e_idx + 2  # +2 because row 0=header, row 1=title
            if row_idx >= n_rows:
                break

            content_cell = table.cell(row_idx, 0)
            content_cell.text = ""
            bg = "FFFFFF" if (e_idx % 2 == 0) else "F8FCFC"
            set_cell_shading(content_cell, bg)

            level = entry["level"]
            text = entry["text"]

            if level == 1:
                _set_generous_padding(content_cell, top=30, bottom=30, left=360, right=100)
                p = content_cell.paragraphs[0]
                bullet = p.add_run("■  ")
                bullet.font.size = Pt(8)
                bullet.font.color.rgb = GOLD
                txt = p.add_run(text)
                txt.font.name = FONT_BODY
                txt.font.size = Pt(9)
                txt.font.color.rgb = DARK_TEXT
            else:
                _set_generous_padding(content_cell, top=20, bottom=20, left=560, right=100)
                p = content_cell.paragraphs[0]
                bullet = p.add_run("•  ")
                bullet.font.size = Pt(8)
                bullet.font.color.rgb = STEEL_GRAY
                txt = p.add_run(text)
                txt.font.name = FONT_BODY
                txt.font.size = Pt(9)
                txt.font.color.rgb = STEEL_GRAY

            # Right column — keep consistent bg
            right_cell = table.cell(row_idx, 1)
            right_cell.text = ""
            set_cell_shading(right_cell, bg)

        # Small spacer between section cards
        spacer = doc.add_paragraph("")
        spacer.paragraph_format.space_before = Pt(2)
        spacer.paragraph_format.space_after = Pt(2)

    # ── Footer legend bar ────────────────────────────────────────────
    _add_toc_legend_bar(doc)


def _collect_toc_entries(plan: SectionPlan) -> list:
    """Collect all subsection/child entries for a section plan as flat list.

    Returns list of dicts: [{"level": 1|2, "text": "..."}]
    """
    entries = []

    if plan.section_type == "segment":
        for seg_name in plan.segment_names:
            entries.append({"level": 1, "text": seg_name})

    elif plan.section_type == "region":
        for region in plan.regions:
            rname = region.get("name", "")
            if rname:
                entries.append({"level": 1, "text": rname})
            for country in region.get("countries", []):
                entries.append({"level": 2, "text": country})

    elif plan.section_type == "competitive":
        _region_labels = {
            "global": "Global Key Players",
            "north_america": "North America Key Players",
            "europe": "Europe Key Players",
            "asia_pacific": "Asia Pacific Key Players",
            "latin_america": "Latin America Key Players",
            "middle_east": "Middle East Key Players",
            "africa": "Africa Key Players",
        }
        for region_key, companies in plan.companies.items():
            if not companies:
                continue
            label = _region_labels.get(region_key,
                                       region_key.replace("_", " ").title() + " Key Players")
            entries.append({"level": 1, "text": label})
            for company in companies:
                entries.append({"level": 2, "text": company})

    else:
        # Generic: overview, key_insights, appendix
        subsections = plan.toc.get("subsections", [])
        for sub in subsections:
            sub_title = _clean_sub_title(sub.get("title", ""))
            if sub_title:
                entries.append({"level": 1, "text": sub_title})
            for child in sub.get("children", []):
                child_title = child.get("title", "")
                if child_title:
                    entries.append({"level": 2, "text": child_title})

    return entries


def _add_toc_legend_bar(doc: Document):
    """Add a styled legend bar at the bottom of the TOC."""
    table = doc.add_table(rows=1, cols=1)
    _remove_table_borders(table)
    _keep_table_on_one_page(table)
    table.alignment = 1

    cell = table.cell(0, 0)
    cell.text = ""
    set_cell_shading(cell, "E8F4F5")
    _set_generous_padding(cell, top=100, bottom=100, left=200, right=200)

    para = cell.paragraphs[0]
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    prefix = para.add_run("ANALYSIS METHODOLOGY:   ")
    prefix.font.name = FONT_LABEL
    prefix.font.size = Pt(8)
    prefix.font.bold = True
    prefix.font.color.rgb = NAVY

    for i, (hex_c, label) in enumerate(_DOT_COLORS):
        if i > 0:
            sep = para.add_run("      ")
            sep.font.size = Pt(8)
        dot = para.add_run("●")
        dot.font.size = Pt(10)
        dot.font.color.rgb = RGBColor(int(hex_c[:2], 16), int(hex_c[2:4], 16), int(hex_c[4:], 16))
        lbl = para.add_run(f" {label} Analysis")
        lbl.font.name = FONT_LABEL
        lbl.font.size = Pt(8)
        lbl.font.color.rgb = DARK_TEXT


# ─── Helpers ──────────────────────────────────────────────────────────────────


def _clean_section_title(title: str) -> str:
    """Clean section title for TOC display.

    Strips year ranges and value/volume suffixes from segment/region titles.
    """
    cleaned = re.sub(r',\s*\d{4}\s*[-–]\s*\d{4}.*$', '', title)
    return cleaned.strip().rstrip(",")


def _clean_sub_title(title: str) -> str:
    """Strip leading 'N.M ' numbering prefix from subsection titles."""
    return re.sub(r'^\d+\.\d+\s+', '', title).strip()
