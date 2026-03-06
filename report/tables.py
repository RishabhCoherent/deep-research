"""
Table builders for python-docx.

4 table types:
  1. Full Forecast Table (13 years + CAGR)
  2. Snapshot Table (3 key years)
  3. YoY Growth Table
  4. Percentage Share Table

All tables use: alternating row shading, bold headers, number formatting.
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

from report.styles import (
    NAVY, GOLD, STEEL_GRAY, WHITE, DARK_TEXT, OFF_WHITE,
    DARK_BLUE, TABLE_ALT_ROW, LIGHT_GRAY, BLACK,
    set_cell_shading, STYLE_CHART_CAPTION,
    FONT_DISPLAY, FONT_BODY, FONT_LABEL,
    _keep_table_on_one_page,
)


# ─── Formatting Helpers ──────────────────────────────────────────────────────


def _fmt_value(val, decimals: int = 2) -> str:
    """Format a numeric value with thousands separator."""
    if val is None:
        return "—"
    try:
        num = float(val)
        return f"{num:,.{decimals}f}"
    except (ValueError, TypeError):
        return str(val)


def _fmt_percent(val) -> str:
    """Format a decimal as percentage: 0.056 -> '5.60%'."""
    if val is None:
        return "—"
    try:
        return f"{float(val) * 100:.2f}%"
    except (ValueError, TypeError):
        return str(val)


def _fmt_cagr(val) -> str:
    """Format CAGR value as percentage."""
    return _fmt_percent(val)


def _style_header_cell(cell):
    """Style a header cell: white text on navy background — PPTX-grade."""
    set_cell_shading(cell, "006B77")
    _add_cell_padding(cell, 6)
    for paragraph in cell.paragraphs:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in paragraph.runs:
            run.font.color.rgb = WHITE
            run.font.bold = True
            run.font.size = Pt(9)
            run.font.name = FONT_BODY


def _add_cell_padding(cell, pt_val: int = 5):
    """Add padding to a table cell for cleaner spacing."""
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn as _qn

    tcPr = cell._element.get_or_add_tcPr()
    tcMar = OxmlElement("w:tcMar")
    for edge in ("top", "bottom", "left", "right"):
        elem = OxmlElement(f"w:{edge}")
        elem.set(_qn("w:w"), str(pt_val * 20))
        elem.set(_qn("w:type"), "dxa")
        tcMar.append(elem)
    tcPr.append(tcMar)


def _style_data_cell(cell, row_idx: int, align: str = "center", bold: bool = False,
                     color: RGBColor = None):
    """Style a data cell with alternating row shading — PPTX-grade.

    Args:
        align: "center", "right", or "left"
        bold: Whether to bold the text
        color: Optional font color override
    """
    # Even rows: subtle blue-tint matching navy theme; odd rows: clean white
    if row_idx % 2 == 0:
        set_cell_shading(cell, "E0F0F0")
    _add_cell_padding(cell, 5)
    alignment_map = {
        "center": WD_ALIGN_PARAGRAPH.CENTER,
        "right": WD_ALIGN_PARAGRAPH.RIGHT,
        "left": WD_ALIGN_PARAGRAPH.LEFT,
    }
    for paragraph in cell.paragraphs:
        paragraph.alignment = alignment_map.get(align, WD_ALIGN_PARAGRAPH.CENTER)
        for run in paragraph.runs:
            run.font.size = Pt(9)
            run.font.name = FONT_BODY
            run.font.bold = bold
            run.font.color.rgb = color if color else DARK_TEXT


def _style_growth_cell(cell, val, row_idx: int):
    """Style a growth/YoY cell with conditional coloring (green ▲ / red ▼) — PPTX-grade."""
    if row_idx % 2 == 0:
        set_cell_shading(cell, "E0F0F0")

    _GREEN = RGBColor(0x00, 0xA6, 0x3E)   # Positive growth
    _RED = RGBColor(0xFB, 0x2C, 0x36)      # Negative / bright red

    try:
        num = float(val)
    except (ValueError, TypeError):
        # Non-numeric — just style normally
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            for run in paragraph.runs:
                run.font.size = Pt(8)
                run.font.name = FONT_BODY
        return

    # Clear cell and rebuild with colored prefix
    cell.text = ""
    para = cell.paragraphs[0]
    para.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    pct = num * 100
    if num > 0:
        prefix_run = para.add_run("\u25b2 ")
        prefix_run.font.color.rgb = _GREEN
        prefix_run.font.size = Pt(7)
        prefix_run.font.name = FONT_BODY
        val_run = para.add_run(f"{pct:.2f}%")
        val_run.font.color.rgb = _GREEN
        val_run.font.size = Pt(8)
        val_run.font.name = FONT_BODY
        val_run.font.bold = True
    elif num < 0:
        prefix_run = para.add_run("\u25bc ")
        prefix_run.font.color.rgb = _RED
        prefix_run.font.size = Pt(7)
        prefix_run.font.name = FONT_BODY
        val_run = para.add_run(f"{pct:.2f}%")
        val_run.font.color.rgb = _RED
        val_run.font.size = Pt(8)
        val_run.font.name = FONT_BODY
        val_run.font.bold = True
    else:
        val_run = para.add_run(f"{pct:.2f}%")
        val_run.font.size = Pt(8)
        val_run.font.name = FONT_BODY
        val_run.font.color.rgb = DARK_TEXT


def _apply_light_borders(table):
    """Apply light professional borders — PPTX-grade: navy top/bottom, subtle inner."""
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else OxmlElement("w:tblPr")

    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        elem = OxmlElement(f"w:{edge}")
        if edge in ("top", "bottom"):
            # Navy accent borders top and bottom
            elem.set(qn("w:val"), "single")
            elem.set(qn("w:sz"), "8")
            elem.set(qn("w:color"), "006B77")
        elif edge in ("insideH", "insideV"):
            # Very subtle inner borders
            elem.set(qn("w:val"), "single")
            elem.set(qn("w:sz"), "4")
            elem.set(qn("w:color"), "E8E8E8")
        else:
            # Subtle left/right outer borders
            elem.set(qn("w:val"), "single")
            elem.set(qn("w:sz"), "4")
            elem.set(qn("w:color"), "E8E8E8")
        elem.set(qn("w:space"), "0")
        borders.append(elem)
    tblPr.append(borders)


def _style_label_cell(cell, row_idx: int, bold: bool = False):
    """Style a row label cell — PPTX-grade."""
    if row_idx % 2 == 0:
        set_cell_shading(cell, "F0F7F7")
    _add_cell_padding(cell, 4)
    for paragraph in cell.paragraphs:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        for run in paragraph.runs:
            run.font.size = Pt(9)
            run.font.name = FONT_BODY
            run.font.bold = bold
            run.font.color.rgb = NAVY if bold else DARK_TEXT


def _style_cagr_header_cell(cell):
    """Style the CAGR header cell with gold background — visually distinct."""
    set_cell_shading(cell, "00BCD4")  # project gold
    _add_cell_padding(cell, 6)
    for paragraph in cell.paragraphs:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in paragraph.runs:
            run.font.color.rgb = WHITE
            run.font.bold = True
            run.font.size = Pt(9)
            run.font.name = FONT_BODY


def _set_row_cant_split(row):
    """Prevent a table row from splitting across pages."""
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    trPr = row._tr.get_or_add_trPr()
    cant_split = OxmlElement("w:cantSplit")
    cant_split.set(qn("w:val"), "1")
    trPr.append(cant_split)


def _set_row_keep_with_next(row):
    """Set keep_with_next on all paragraphs in a row.

    Forces Word to keep this row on the same page as the following row.
    Use on the header row so it never gets stranded at the bottom of a page.
    """
    for cell in row.cells:
        for para in cell.paragraphs:
            para.paragraph_format.keep_with_next = True


def _set_caption_keep_with_next(doc, text: str) -> None:
    """Add a caption paragraph that stays with the table that follows it."""
    para = doc.add_paragraph(text, style=STYLE_CHART_CAPTION)
    para.paragraph_format.keep_with_next = True
    return para


def _set_table_width(table, width_inches: float = 9.5):
    """Set consistent table formatting."""
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True


# ─── Table Builders ──────────────────────────────────────────────────────────


def add_forecast_table(
    doc: Document,
    title: str,
    items: dict,
    years: list[str],
    cagr_key: str,
    unit: str = "",
    value_decimals: int = 2,
    forecast_start_yr: str = None,
):
    """Add a full forecast table with all years + CAGR.

    Args:
        doc: Document to add to.
        title: Table title (displayed above).
        items: {"Item Name": {"forecast": {"2020": x, ...}, "cagr_...": x}}
        years: List of year strings.
        cagr_key: The CAGR dict key (e.g. "cagr_2025_2032").
        unit: Unit string to display in title.
        value_decimals: Decimal places for values.
        forecast_start_yr: First forecast year. When set, shades forecast year
                           headers in a lighter navy to distinguish from historical.
    """
    if not items or not years:
        return

    unit_str = f" ({unit})" if unit else ""
    _set_caption_keep_with_next(doc, f"{title}{unit_str}")

    # Columns: Label + years + CAGR
    col_headers = [""] + years + [cagr_key.replace("_", " ").upper()]
    n_cols = len(col_headers)
    n_rows = 1 + len(items)

    table = doc.add_table(rows=n_rows, cols=n_cols)
    table.style = "Table Grid"
    _keep_table_on_one_page(table)
    _set_table_width(table)
    _apply_light_borders(table)

    # Header row — forecast years lighter navy, CAGR column gold
    _set_row_cant_split(table.rows[0])
    _set_row_keep_with_next(table.rows[0])
    for j, header in enumerate(col_headers):
        cell = table.cell(0, j)
        cell.text = header
        if j == n_cols - 1:
            _style_cagr_header_cell(cell)
        elif forecast_start_yr and j > 0 and j <= len(years):
            yr = years[j - 1]
            if str(yr) >= str(forecast_start_yr):
                set_cell_shading(cell, "008B96")
                _add_cell_padding(cell, 6)
                for paragraph in cell.paragraphs:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    for run in paragraph.runs:
                        run.font.color.rgb = WHITE
                        run.font.bold = True
                        run.font.size = Pt(9)
                        run.font.name = FONT_BODY
                continue
            else:
                _style_header_cell(cell)
        else:
            _style_header_cell(cell)

    # Data rows
    for i, (name, data) in enumerate(items.items()):
        row_idx = i + 1
        forecast = data.get("forecast", {})
        cagr_val = data.get(cagr_key)

        _set_row_cant_split(table.rows[row_idx])

        # Label
        cell = table.cell(row_idx, 0)
        cell.text = name
        _style_label_cell(cell, i, bold=True)

        # Year values (right-aligned)
        for j, year in enumerate(years):
            cell = table.cell(row_idx, j + 1)
            cell.text = _fmt_value(forecast.get(year), value_decimals)
            _style_data_cell(cell, i, align="right")

        # CAGR — bold navy to match gold header
        cell = table.cell(row_idx, n_cols - 1)
        cell.text = _fmt_cagr(cagr_val)
        _style_data_cell(cell, i, align="right", bold=True, color=NAVY)

    doc.add_paragraph("")  # spacing


def add_snapshot_table(
    doc: Document,
    title: str,
    items: dict,
    snapshot_years: list[str],
    cagr_key: str,
    unit: str = "",
    forecast_start_yr: str = None,
):
    """Add a compact snapshot table with 3 key years + CAGR.

    Same structure as forecast table but fewer columns.

    Args:
        forecast_start_yr: When set, shades forecast year headers in lighter navy.
    """
    if not items or not snapshot_years:
        return

    unit_str = f" ({unit})" if unit else ""
    _set_caption_keep_with_next(doc, f"{title}{unit_str}")

    col_headers = [""] + snapshot_years + ["CAGR"]
    n_cols = len(col_headers)
    n_rows = 1 + len(items)

    table = doc.add_table(rows=n_rows, cols=n_cols)
    table.style = "Table Grid"
    _keep_table_on_one_page(table)
    _set_table_width(table)
    _apply_light_borders(table)

    # Header row — keep with first data row
    _set_row_cant_split(table.rows[0])
    _set_row_keep_with_next(table.rows[0])
    for j, header in enumerate(col_headers):
        cell = table.cell(0, j)
        cell.text = header
        if j == n_cols - 1:
            # CAGR column: gold header
            _style_cagr_header_cell(cell)
        elif forecast_start_yr and j > 0 and j <= len(snapshot_years):
            yr = snapshot_years[j - 1]
            if str(yr) >= str(forecast_start_yr):
                set_cell_shading(cell, "008B96")
                _add_cell_padding(cell, 6)
                for paragraph in cell.paragraphs:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    for run in paragraph.runs:
                        run.font.color.rgb = WHITE
                        run.font.bold = True
                        run.font.size = Pt(9)
                        run.font.name = FONT_BODY
                continue
            else:
                _style_header_cell(cell)
        else:
            _style_header_cell(cell)

    for i, (name, data) in enumerate(items.items()):
        row_idx = i + 1
        forecast = data.get("forecast", {})
        cagr_val = data.get(cagr_key)

        _set_row_cant_split(table.rows[row_idx])

        cell = table.cell(row_idx, 0)
        cell.text = name
        _style_label_cell(cell, i, bold=True)

        for j, year in enumerate(snapshot_years):
            cell = table.cell(row_idx, j + 1)
            cell.text = _fmt_value(forecast.get(year))
            _style_data_cell(cell, i, align="right")

        # CAGR cell: bold navy to match gold header
        cell = table.cell(row_idx, n_cols - 1)
        cell.text = _fmt_cagr(cagr_val)
        _style_data_cell(cell, i, align="right", bold=True, color=NAVY)

    doc.add_paragraph("")


def add_yoy_table(
    doc: Document,
    title: str,
    items: dict,
    years: list[str] | None = None,
):
    """Add a YoY growth table with percentage values.

    Args:
        items: {"Item": {"yoy_growth": {"2021": 0.05, ...}}}
        years: Optional year filter. If None, discovers from data.
    """
    if not items:
        return

    # Discover years from data if not provided
    if years is None:
        all_years = set()
        for data in items.values():
            all_years.update(data.get("yoy_growth", {}).keys())
        years = sorted(all_years, key=int)

    if not years:
        return

    _set_caption_keep_with_next(doc, f"{title} — Year-over-Year Growth")

    col_headers = [""] + years
    n_cols = len(col_headers)
    n_rows = 1 + len(items)

    table = doc.add_table(rows=n_rows, cols=n_cols)
    table.style = "Table Grid"
    _keep_table_on_one_page(table)
    _set_table_width(table)
    _apply_light_borders(table)

    _set_row_cant_split(table.rows[0])
    _set_row_keep_with_next(table.rows[0])
    for j, header in enumerate(col_headers):
        cell = table.cell(0, j)
        cell.text = header
        _style_header_cell(cell)

    for i, (name, data) in enumerate(items.items()):
        row_idx = i + 1
        yoy = data.get("yoy_growth", {})

        _set_row_cant_split(table.rows[row_idx])

        cell = table.cell(row_idx, 0)
        cell.text = name
        _style_label_cell(cell, i, bold=True)

        for j, year in enumerate(years):
            cell = table.cell(row_idx, j + 1)
            raw_val = yoy.get(year)
            _style_growth_cell(cell, raw_val, i)

    doc.add_paragraph("")


def add_pct_share_table(
    doc: Document,
    title: str,
    items: dict,
    years: list[str] | None = None,
):
    """Add a percentage share table.

    Args:
        items: {"Item": {"percentage_share": {"2020": 0.43, ...}}}
        years: Optional year filter.
    """
    if not items:
        return

    if years is None:
        all_years = set()
        for data in items.values():
            all_years.update(data.get("percentage_share", {}).keys())
        years = sorted(all_years, key=int)

    if not years:
        return

    _set_caption_keep_with_next(doc, f"{title} — Market Share (%)")

    col_headers = [""] + years
    n_cols = len(col_headers)
    n_rows = 1 + len(items)

    table = doc.add_table(rows=n_rows, cols=n_cols)
    table.style = "Table Grid"
    _keep_table_on_one_page(table)
    _set_table_width(table)
    _apply_light_borders(table)

    _set_row_cant_split(table.rows[0])
    _set_row_keep_with_next(table.rows[0])
    for j, header in enumerate(col_headers):
        cell = table.cell(0, j)
        cell.text = header
        _style_header_cell(cell)

    for i, (name, data) in enumerate(items.items()):
        row_idx = i + 1
        pct = data.get("percentage_share", {})

        _set_row_cant_split(table.rows[row_idx])

        cell = table.cell(row_idx, 0)
        cell.text = name
        _style_label_cell(cell, i, bold=True)

        for j, year in enumerate(years):
            cell = table.cell(row_idx, j + 1)
            cell.text = _fmt_percent(pct.get(year))
            _style_data_cell(cell, i, align="right")

    doc.add_paragraph("")
