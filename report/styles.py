"""
Document styling — fonts, colors, page setup, custom styles.

Creates a landscape 11"×8.5" document with PPTX-grade corporate styling.
Design system: Dark Teal (#006B77) + Bright Cyan (#00BCD4) + Slate (#5A7D8C)
All style names are defined as constants for reuse across modules.
"""

import io

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.section import WD_ORIENT, WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ─── Color Palette (PPTX Design System) ─────────────────────────────────────

# Primary colors
NAVY = RGBColor(0x00, 0x6B, 0x77)          # Primary dark — panels, headers, backgrounds
GOLD = RGBColor(0x00, 0xBC, 0xD4)          # Bright accent — KPI values, accent bars, labels
STEEL_GRAY = RGBColor(0x5A, 0x7D, 0x8C)    # Tertiary — third-tier panels, descriptions
OFF_WHITE = RGBColor(0xE8, 0xF4, 0xF5)     # Muted text, subtle backgrounds
DARK_TEXT = RGBColor(0x00, 0x3F, 0x48)      # Body text on light backgrounds
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
BLACK = RGBColor(0x00, 0x00, 0x00)

# Green accent colors (mixed into teal palette for variety)
TEAL_GREEN = RGBColor(0x00, 0x96, 0x88)    # Green-teal accent (#009688)
FOREST_GREEN = RGBColor(0x2E, 0x7D, 0x32)  # Deep green accent (#2E7D32)
LIGHT_TEAL_GREEN = RGBColor(0x4D, 0xB6, 0xAC)  # Soft green-teal (#4DB6AC)

# Semantic indicator colors
BRIGHT_RED = RGBColor(0xFB, 0x2C, 0x36)    # High risk / strong decline
MINT_GREEN = RGBColor(0x7B, 0xF1, 0xA8)    # Positive growth
CORAL_RED = RGBColor(0xFF, 0xA2, 0xA2)     # Moderate decline
ORANGE_RISK = RGBColor(0xFF, 0x69, 0x00)   # Medium risk
AMBER = RGBColor(0xF0, 0xB1, 0x00)         # Structural risk / warning

# Table-specific
TABLE_HEADER_BG = NAVY
TABLE_ALT_ROW = RGBColor(0xF0, 0xF7, 0xF7)  # Very subtle alternation

# Hex versions for matplotlib
HEX_NAVY = "#006B77"
HEX_GOLD = "#00BCD4"
HEX_STEEL_GRAY = "#5A7D8C"
HEX_OFF_WHITE = "#E8F4F5"
HEX_DARK_TEXT = "#003F48"
HEX_BRIGHT_RED = "#FB2C36"
HEX_MINT_GREEN = "#7BF1A8"
HEX_CORAL_RED = "#FFA2A2"
HEX_ORANGE = "#FF6900"
HEX_AMBER = "#F0B100"
HEX_TEAL_GREEN = "#009688"
HEX_FOREST_GREEN = "#2E7D32"
HEX_LIGHT_TEAL_GREEN = "#4DB6AC"

# Chart palette (teal + green + cyan system)
CHART_COLORS = ["#006B77", "#009688", "#00BCD4", "#2E7D32", "#5A7D8C", "#4DB6AC"]

# Backward-compatibility aliases (so existing imports don't break)
DARK_BLUE = NAVY
ACCENT_BLUE = NAVY
ACCENT_ORANGE = GOLD
LIGHT_GRAY = OFF_WHITE
MID_GRAY = STEEL_GRAY
HEX_DARK_BLUE = HEX_NAVY
HEX_ACCENT_BLUE = HEX_NAVY
HEX_ACCENT_ORANGE = HEX_GOLD
HEX_LIGHT_GRAY = HEX_OFF_WHITE
TABLE_ALT_ROW_LEGACY = RGBColor(0xD8, 0xED, 0xED)

# ─── Font System ─────────────────────────────────────────────────────────────

FONT_DISPLAY = "Georgia"        # Serif: titles, KPI numbers (Oranienbaum substitute)
FONT_BODY = "Calibri"           # Sans: body text (Quattrocento Sans equivalent)
FONT_LABEL = "Calibri Light"    # Utility: section labels, captions (Liter equivalent)

# ─── Style Names ─────────────────────────────────────────────────────────────

STYLE_REPORT_TITLE = "Report Title"
STYLE_SECTION_HEADING = "Section Heading"
STYLE_SUBSECTION = "Subsection Heading"
STYLE_SUB_SUBSECTION = "Sub-Subsection Heading"
STYLE_BODY = "Report Body"
STYLE_CHART_CAPTION = "Chart Caption"
STYLE_TABLE_HEADER = "Table Header Cell"
STYLE_COVER_SUBTITLE = "Cover Subtitle"
STYLE_BIBLIOGRAPHY = "Bibliography Entry"
STYLE_TOC_SECTION = "TOC Section Entry"
STYLE_TOC_SUBSECTION = "TOC Subsection Entry"
STYLE_TOC_CHILD = "TOC Child Entry"
STYLE_SECTION_LABEL = "Section Label"


# ─── Report Title Registry (for slide headers/footers) ──────────────────────

_report_title = ""


def set_report_title(title: str):
    """Store report title for use in slide headers/footers."""
    global _report_title
    _report_title = title


# ─── Document Creation ───────────────────────────────────────────────────────


def create_styled_document() -> Document:
    """Create a landscape document with all custom styles registered."""
    doc = Document()

    # Page setup: landscape 11" × 8.5" with 0.75" margins
    section = doc.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width = Inches(11)
    section.page_height = Inches(8.5)
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)
    section.top_margin = Inches(0.6)
    section.bottom_margin = Inches(0.5)

    # Default font
    style = doc.styles["Normal"]
    font = style.font
    font.name = FONT_BODY
    font.size = Pt(10)
    font.color.rgb = DARK_TEXT

    _register_styles(doc)

    return doc


def _register_styles(doc: Document):
    """Register all custom paragraph/character styles."""
    styles = doc.styles

    # Report Title (cover page) — Georgia serif, large
    s = styles.add_style(STYLE_REPORT_TITLE, 1)
    s.font.name = FONT_DISPLAY
    s.font.size = Pt(36)
    s.font.bold = True
    s.font.color.rgb = NAVY
    s.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    s.paragraph_format.space_after = Pt(8)

    # Cover Subtitle — Calibri Light, gold
    s = styles.add_style(STYLE_COVER_SUBTITLE, 1)
    s.font.name = FONT_LABEL
    s.font.size = Pt(16)
    s.font.color.rgb = GOLD
    s.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    s.paragraph_format.space_after = Pt(6)

    # Section Heading — Georgia serif, white on navy band
    s = styles.add_style(STYLE_SECTION_HEADING, 1)
    s.font.name = FONT_DISPLAY
    s.font.size = Pt(28)
    s.font.bold = True
    s.font.color.rgb = WHITE
    s.paragraph_format.space_before = Pt(28)
    s.paragraph_format.space_after = Pt(18)
    s.paragraph_format.keep_with_next = True

    # Subsection Heading — Georgia serif, dark text
    s = styles.add_style(STYLE_SUBSECTION, 1)
    s.font.name = FONT_DISPLAY
    s.font.size = Pt(17)
    s.font.bold = True
    s.font.color.rgb = DARK_TEXT
    s.paragraph_format.space_before = Pt(20)
    s.paragraph_format.space_after = Pt(8)
    s.paragraph_format.keep_with_next = True

    # Sub-Subsection Heading — Calibri, dark text
    s = styles.add_style(STYLE_SUB_SUBSECTION, 1)
    s.font.name = FONT_BODY
    s.font.size = Pt(13)
    s.font.bold = True
    s.font.color.rgb = DARK_TEXT
    s.paragraph_format.space_before = Pt(12)
    s.paragraph_format.space_after = Pt(4)
    s.paragraph_format.keep_with_next = True
    s.paragraph_format.left_indent = Cm(0.3)

    # Body Text — Calibri, dark text (presentation-friendly spacing)
    s = styles.add_style(STYLE_BODY, 1)
    s.font.name = FONT_BODY
    s.font.size = Pt(11)
    s.font.color.rgb = DARK_TEXT
    s.paragraph_format.space_after = Pt(6)
    s.paragraph_format.line_spacing = Pt(17)

    # Chart Caption — Calibri Light, steel gray
    s = styles.add_style(STYLE_CHART_CAPTION, 1)
    s.font.name = FONT_LABEL
    s.font.size = Pt(10)
    s.font.bold = True
    s.font.color.rgb = STEEL_GRAY
    s.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    s.paragraph_format.space_before = Pt(6)
    s.paragraph_format.space_after = Pt(10)

    # Section Label — gold ALL-CAPS above headings
    s = styles.add_style(STYLE_SECTION_LABEL, 1)
    s.font.name = FONT_LABEL
    s.font.size = Pt(9)
    s.font.bold = True
    s.font.color.rgb = GOLD
    s.paragraph_format.space_before = Pt(4)
    s.paragraph_format.space_after = Pt(2)
    s.paragraph_format.keep_with_next = True

    # Bibliography Entry
    s = styles.add_style(STYLE_BIBLIOGRAPHY, 1)
    s.font.name = FONT_BODY
    s.font.size = Pt(9)
    s.font.color.rgb = DARK_TEXT
    s.paragraph_format.space_after = Pt(3)
    s.paragraph_format.left_indent = Cm(1)
    s.paragraph_format.first_line_indent = Cm(-1)

    # TOC Section Entry (level 0 — Georgia, navy)
    s = styles.add_style(STYLE_TOC_SECTION, 1)
    s.font.name = FONT_DISPLAY
    s.font.size = Pt(11)
    s.font.bold = True
    s.font.color.rgb = NAVY
    s.paragraph_format.space_before = Pt(8)
    s.paragraph_format.space_after = Pt(2)
    s.paragraph_format.left_indent = Inches(0.3)

    # TOC Subsection Entry (level 1 — navy)
    s = styles.add_style(STYLE_TOC_SUBSECTION, 1)
    s.font.name = FONT_BODY
    s.font.size = Pt(10)
    s.font.color.rgb = NAVY
    s.paragraph_format.space_after = Pt(1)
    s.paragraph_format.left_indent = Inches(0.7)

    # TOC Child Entry (level 2 — steel gray)
    s = styles.add_style(STYLE_TOC_CHILD, 1)
    s.font.name = FONT_BODY
    s.font.size = Pt(9)
    s.font.color.rgb = STEEL_GRAY
    s.paragraph_format.space_after = Pt(0)
    s.paragraph_format.left_indent = Inches(1.1)


# ─── Layout Helpers (private) ────────────────────────────────────────────────


def _remove_table_borders(table):
    """Remove all borders from a table — for layout-only tables."""
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else OxmlElement("w:tblPr")
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        elem = OxmlElement(f"w:{edge}")
        elem.set(qn("w:val"), "none")
        elem.set(qn("w:sz"), "0")
        elem.set(qn("w:space"), "0")
        elem.set(qn("w:color"), "FFFFFF")
        borders.append(elem)
    tblPr.append(borders)


def _keep_table_on_one_page(table):
    """Prevent a table from splitting across pages.

    Sets two Word XML properties on every row:
      • w:cantSplit  – prevents a single row from being split across a page break
      • w:keepNext   – (on all rows except the last) tells Word to keep consecutive
                        rows on the same page, effectively gluing the whole table together
    """
    rows = table.rows
    for idx, row in enumerate(rows):
        tr = row._tr
        trPr = tr.get_or_add_trPr()
        # Prevent individual row from splitting
        cant_split = OxmlElement("w:cantSplit")
        trPr.append(cant_split)
        # Keep this row with the next one (skip on last row)
        if idx < len(rows) - 1:
            for cell in row.cells:
                for para in cell.paragraphs:
                    pPr = para._p.get_or_add_pPr()
                    keep_next = OxmlElement("w:keepNext")
                    pPr.append(keep_next)


def _set_generous_padding(cell, top=200, bottom=200, left=240, right=240):
    """Set generous padding on a table cell (in dxa units; 200 dxa ≈ 10pt)."""
    tcPr = cell._element.get_or_add_tcPr()
    tcMar = OxmlElement("w:tcMar")
    for edge, val in [("top", top), ("bottom", bottom), ("left", left), ("right", right)]:
        elem = OxmlElement(f"w:{edge}")
        elem.set(qn("w:w"), str(val))
        elem.set(qn("w:type"), "dxa")
        tcMar.append(elem)
    tcPr.append(tcMar)


def _add_left_accent_border(cell, color_hex: str, thickness: str = "36"):
    """Add only a left accent border to a cell, no other borders."""
    tcPr = cell._element.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), thickness)
    left.set(qn("w:space"), "0")
    left.set(qn("w:color"), color_hex)
    tcBorders.append(left)
    for edge in ("top", "right", "bottom"):
        e = OxmlElement(f"w:{edge}")
        e.set(qn("w:val"), "none")
        e.set(qn("w:sz"), "0")
        e.set(qn("w:space"), "0")
        e.set(qn("w:color"), "FFFFFF")
        tcBorders.append(e)
    tcPr.append(tcBorders)


# ─── Heading & Text Helpers ──────────────────────────────────────────────────


def add_section_label(doc: Document, label: str):
    """Add a gold ALL-CAPS section label above headings (PPTX pattern).

    Creates the two-line heading pattern: small gold label + large heading below.
    """
    para = doc.add_paragraph(style=STYLE_SECTION_LABEL)
    run = para.add_run(label.upper())
    run.font.name = FONT_LABEL
    run.font.size = Pt(9)
    run.font.bold = True
    run.font.color.rgb = GOLD
    # Add letter-spacing via XML
    rPr = run._element.get_or_add_rPr()
    spacing = OxmlElement("w:spacing")
    spacing.set(qn("w:val"), "40")  # 2pt tracking
    rPr.append(spacing)


# ─── Slide Layout System (PPTX-style repeating header/footer) ────────────────


def _gradient_colors(start_hex, end_hex, steps):
    """Generate a list of hex color strings forming a smooth gradient."""
    s = [int(start_hex[i:i+2], 16) for i in (0, 2, 4)]
    e = [int(end_hex[i:i+2], 16) for i in (0, 2, 4)]
    colors = []
    for i in range(steps):
        t = i / max(steps - 1, 1)
        r = int(s[0] + (e[0] - s[0]) * t)
        g = int(s[1] + (e[1] - s[1]) * t)
        b = int(s[2] + (e[2] - s[2]) * t)
        colors.append(f"{r:02X}{g:02X}{b:02X}")
    return colors


def _set_cell_width_xml(cell, width_dxa):
    """Force cell width via XML w:tcW element."""
    tcPr = cell._element.get_or_add_tcPr()
    existing = tcPr.find(qn("w:tcW"))
    if existing is not None:
        tcPr.remove(existing)
    tcW = OxmlElement("w:tcW")
    tcW.set(qn("w:w"), str(width_dxa))
    tcW.set(qn("w:type"), "dxa")
    tcPr.append(tcW)


def _set_table_grid(table, col_widths_dxa):
    """Set explicit w:tblGrid column widths so Word honors exact sizing."""
    tbl = table._tbl
    existing = tbl.find(qn("w:tblGrid"))
    if existing is not None:
        tbl.remove(existing)
    grid = OxmlElement("w:tblGrid")
    for w in col_widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(w))
        grid.append(col)
    tblPr = tbl.tblPr
    if tblPr is not None:
        tblPr.addnext(grid)
    else:
        tbl.insert(0, grid)


def _set_table_full_bleed(table, left_margin_dxa, page_width_dxa=15840):
    """Make a header/footer table bleed edge-to-edge past page margins.

    Sets table width, negative indent, fixed layout, and zero cell spacing
    so the table covers the full page width with no gaps.
    """
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else OxmlElement("w:tblPr")

    # Force table width to full page width (11" = 15840 dxa)
    existing_tblW = tblPr.find(qn("w:tblW"))
    if existing_tblW is not None:
        tblPr.remove(existing_tblW)
    tblW = OxmlElement("w:tblW")
    tblW.set(qn("w:w"), str(page_width_dxa))
    tblW.set(qn("w:type"), "dxa")
    tblPr.append(tblW)

    # Negative indent to pull the table left past the left margin
    existing_tblInd = tblPr.find(qn("w:tblInd"))
    if existing_tblInd is not None:
        tblPr.remove(existing_tblInd)
    tblInd = OxmlElement("w:tblInd")
    tblInd.set(qn("w:w"), str(-left_margin_dxa))
    tblInd.set(qn("w:type"), "dxa")
    tblPr.append(tblInd)

    # Fixed layout prevents Word from auto-shrinking columns
    existing_layout = tblPr.find(qn("w:tblLayout"))
    if existing_layout is not None:
        tblPr.remove(existing_layout)
    tblLayout = OxmlElement("w:tblLayout")
    tblLayout.set(qn("w:type"), "fixed")
    tblPr.append(tblLayout)

    # Zero cell spacing — prevents white gaps between gradient cells
    existing_spacing = tblPr.find(qn("w:tblCellSpacing"))
    if existing_spacing is not None:
        tblPr.remove(existing_spacing)
    cellSpacing = OxmlElement("w:tblCellSpacing")
    cellSpacing.set(qn("w:w"), "0")
    cellSpacing.set(qn("w:type"), "dxa")
    tblPr.append(cellSpacing)


def start_slide_section(doc, section_title, report_title=""):
    """Start a new Word section with PPTX-style gradient header/footer bars.

    Creates a new page with a gradient header (teal fading to white) and a
    thin gradient footer that auto-repeat on every page within this section.
    """
    new_section = doc.add_section(WD_SECTION.NEW_PAGE)

    # Carry over page dimensions from section 0
    new_section.orientation = WD_ORIENT.LANDSCAPE
    new_section.page_width = Inches(11)
    new_section.page_height = Inches(8.5)
    new_section.left_margin = Inches(0.75)
    new_section.right_margin = Inches(0.75)
    new_section.top_margin = Inches(1.35)    # room for tall header
    new_section.bottom_margin = Inches(0.6)

    # Push header/footer flush to page edge
    new_section.header_distance = Inches(0.0)
    new_section.footer_distance = Inches(0.0)

    _build_slide_header(new_section, section_title, report_title)
    _build_slide_footer(new_section, report_title)

    return new_section


def _build_slide_header(section, section_title, report_title):
    """Build a PPTX-style header with smooth teal→white gradient (20 columns)."""
    header = section.header
    header.is_linked_to_previous = False

    for p in header.paragraphs:
        p._element.getparent().remove(p._element)

    page_width_dxa = 15840  # 11"
    margin_dxa = 1080       # 0.75"

    # 20 narrow columns for a seamless gradient: teal → white
    n_cols = 20
    gradient = _gradient_colors("339CA5", "FFFFFF", n_cols)

    # Wide title col + 18 narrow gradient cols + wide end col = 15840
    title_w = 3960   # ~2.75" for section title text
    end_w = 2160     # ~1.5" for report title text
    mid_w = 540      # each gradient spacer (~0.375")
    col_widths = [title_w] + [mid_w] * 18 + [end_w]
    # 3960 + 18×540 + 2160 = 15840 ✓

    header_table = header.add_table(rows=1, cols=n_cols, width=Inches(11))
    _remove_table_borders(header_table)
    _set_table_full_bleed(header_table, margin_dxa, page_width_dxa)
    _set_table_grid(header_table, col_widths)

    for i in range(n_cols):
        cell = header_table.cell(0, i)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        set_cell_shading(cell, gradient[i])
        _set_cell_width_xml(cell, col_widths[i])
        pad_l = 600 if i == 0 else 0
        pad_r = 200 if i == n_cols - 1 else 0
        _set_generous_padding(cell, top=280, bottom=280, left=pad_l, right=pad_r)
        cell.paragraphs[0].paragraph_format.space_before = Pt(0)
        cell.paragraphs[0].paragraph_format.space_after = Pt(0)

    # Section title (first cell — darkest teal, white text)
    left_para = header_table.cell(0, 0).paragraphs[0]
    run = left_para.add_run(section_title.upper())
    run.font.name = FONT_LABEL
    run.font.size = Pt(16)
    run.font.bold = True
    run.font.color.rgb = WHITE

    # Report title (last cell — white bg, navy text)
    right_para = header_table.cell(0, n_cols - 1).paragraphs[0]
    right_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = right_para.add_run(report_title[:40] if report_title else "")
    run.font.name = FONT_LABEL
    run.font.size = Pt(9)
    run.font.color.rgb = NAVY


def _build_slide_footer(section, report_title):
    """Build a thin footer with smooth gradient (16 columns), source + page only."""
    footer = section.footer
    footer.is_linked_to_previous = False

    for p in footer.paragraphs:
        p._element.getparent().remove(p._element)

    margin_dxa = 1080
    page_width_dxa = 15840

    # 16 narrow columns for seamless gradient
    n_cols = 16
    gradient = _gradient_colors("66BCC3", "FFFFFF", n_cols)

    # Source col + 14 gradient cols + page number col = 15840
    source_w = 4320
    page_w = 1440
    mid_w = 720
    col_widths = [source_w] + [mid_w] * 14 + [page_w]
    # 4320 + 14×720 + 1440 = 15840 ✓

    footer_table = footer.add_table(rows=1, cols=n_cols, width=Inches(11))
    _remove_table_borders(footer_table)
    _set_table_full_bleed(footer_table, margin_dxa, page_width_dxa)
    _set_table_grid(footer_table, col_widths)

    for i in range(n_cols):
        cell = footer_table.cell(0, i)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        set_cell_shading(cell, gradient[i])
        _set_cell_width_xml(cell, col_widths[i])
        pad_l = 200 if i == 0 else 0
        pad_r = 200 if i == n_cols - 1 else 0
        _set_generous_padding(cell, top=80, bottom=80, left=pad_l, right=pad_r)
        cell.paragraphs[0].paragraph_format.space_before = Pt(0)
        cell.paragraphs[0].paragraph_format.space_after = Pt(0)

    # Source label (col 0 — darkest, white text)
    lp = footer_table.cell(0, 0).paragraphs[0]
    lr = lp.add_run("Source: Primary Research, Secondary Research, Databases")
    lr.font.name = FONT_LABEL
    lr.font.size = Pt(7)
    lr.font.color.rgb = WHITE
    lr.font.italic = True

    # Page number (last col — lightest, dark text)
    rp = footer_table.cell(0, n_cols - 1).paragraphs[0]
    rp.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    page_run = rp.add_run("Page ")
    page_run.font.name = FONT_LABEL
    page_run.font.size = Pt(7)
    page_run.font.color.rgb = DARK_TEXT
    page_run.font.bold = True

    # Auto page number field
    fldChar_begin = OxmlElement("w:fldChar")
    fldChar_begin.set(qn("w:fldCharType"), "begin")
    num_run = rp.add_run()
    num_run.font.name = FONT_LABEL
    num_run.font.size = Pt(7)
    num_run.font.color.rgb = DARK_TEXT
    num_run.font.bold = True
    num_run._element.append(fldChar_begin)

    instrText = OxmlElement("w:instrText")
    instrText.set(qn("xml:space"), "preserve")
    instrText.text = " PAGE "
    num_run._element.append(instrText)

    fldChar_end = OxmlElement("w:fldChar")
    fldChar_end.set(qn("w:fldCharType"), "end")
    num_run._element.append(fldChar_end)


def add_slide_break(doc):
    """Add a page break to start a new slide within the current section.

    The current section's header/footer bars continue on the new page.
    """
    doc.add_page_break()


# ─── Section Heading (creates new Word section with header/footer) ───────────


def add_section_heading(doc: Document, text: str, section_label: str = ""):
    """Add a section heading — creates a new Word section with slide header/footer.

    Only creates the header/footer bars for the new section page. The body-level
    title is rendered separately by add_section_intro_strip() for a cleaner
    presentation layout (header title + one visual body title = 2 total).
    """
    start_slide_section(doc, text, _report_title)


def add_subsection_heading(doc: Document, text: str):
    """Add a subsection heading with a thick gold left accent border."""
    para = doc.add_paragraph(text, style=STYLE_SUBSECTION)

    # Add thick left border (gold accent bar)
    pPr = para._element.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), "48")       # 6pt thick
    left.set(qn("w:space"), "8")
    left.set(qn("w:color"), "00BCD4")  # Gold
    pBdr.append(left)
    pPr.append(pBdr)

    # Warm off-white background
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), "E8F4F5")
    shd.set(qn("w:val"), "clear")
    pPr.append(shd)


def add_sub_subsection_heading(doc: Document, text: str):
    """Add a sub-subsection heading with a thin steel gray left border."""
    para = doc.add_paragraph(text, style=STYLE_SUB_SUBSECTION)

    # Thin left border (steel gray accent)
    pPr = para._element.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), "18")       # 2.25pt
    left.set(qn("w:space"), "6")
    left.set(qn("w:color"), "5A7D8C")  # Steel gray
    pBdr.append(left)
    pPr.append(pBdr)


def add_body_text(doc: Document, text: str):
    """Add body text paragraph."""
    doc.add_paragraph(text, style=STYLE_BODY)


# ─── Citation Registry ───────────────────────────────────────────────────────

_citation_registry: dict[str, dict] = {}


def set_citation_registry(citations: list[dict]):
    """Set the citation registry for resolving inline [src_xxx_nnn] to clickable hyperlinks."""
    global _citation_registry
    _citation_registry = {}
    for c in citations:
        _citation_registry[c["id"]] = {
            "title": c.get("title", ""),
            "url": c.get("url", ""),
            "publisher": c.get("publisher", ""),
        }


# ─── Shared Markdown Renderer ────────────────────────────────────────────────


def clean_markdown(text: str) -> str:
    """Strip all markdown bold/italic markers from text, keeping content."""
    import re
    if not text:
        return text
    text = re.sub(r'\*\*(.*?)\*\*', r'\1', text)
    text = re.sub(r'\*(.*?)\*', r'\1', text)
    return text.strip()


def add_formatted_text(doc: Document, text: str):
    """Render markdown-formatted text to docx with proper bold/italic/table support.

    This is the ONE shared renderer for the entire project.
    """
    import re as _re

    if not text:
        return

    # Pre-process: split mixed blocks (text + table in one \n\n block) into
    # separate blocks so tables are detected correctly.
    raw_blocks = text.split("\n\n")
    blocks = []
    for block in raw_blocks:
        block = block.strip()
        if not block:
            continue
        # If block contains pipe-table lines mixed with non-table lines,
        # split into text part + table part
        lines = block.split("\n")
        table_start = None
        for li, line in enumerate(lines):
            stripped = line.strip()
            if stripped.startswith("|") and "|" in stripped[1:]:
                table_start = li
                break
        if table_start and table_start > 0:
            # Lines before the table → text block
            pre_text = "\n".join(lines[:table_start]).strip()
            if pre_text:
                blocks.append(pre_text)
            # Table lines → table block
            tbl_text = "\n".join(lines[table_start:]).strip()
            if tbl_text:
                blocks.append(tbl_text)
        else:
            blocks.append(block)

    for para_text in blocks:
        para_text = para_text.strip()
        if not para_text:
            continue

        # Headings
        if para_text.startswith("###"):
            header = para_text.lstrip("#").strip()
            add_sub_subsection_heading(doc, clean_markdown(header))
            continue
        if para_text.startswith("##"):
            header = para_text.lstrip("#").strip()
            add_subsection_heading(doc, clean_markdown(header))
            continue
        if para_text.startswith("#"):
            continue

        # Markdown tables — detect by checking if ANY line starts with |
        lines = para_text.split("\n")
        table_lines = [l for l in lines if l.strip().startswith("|")]
        if len(table_lines) >= 2:
            add_markdown_table(doc, para_text)
            continue

        # Bullet lists (lines starting with - , * , or • )
        non_empty = [l for l in lines if l.strip()]
        if non_empty and all(l.strip().startswith(("- ", "* ", "• ")) for l in non_empty):
            for line in non_empty:
                line = _re.sub(r'^\s*[-•]\s+', '', line)
                line = _re.sub(r'^\s*\*\s+', '', line)
                line = line.strip()
                if line:
                    _render_bullet_paragraph(doc, line)
            continue

        # Mixed block: heading lines + bullet lines in one block
        # Render line by line
        has_headings = any(l.strip().startswith("#") for l in lines)
        has_bullets = any(l.strip().startswith(("- ", "* ", "• ")) for l in lines)
        if has_headings or (has_bullets and not all(
                l.strip().startswith(("- ", "* ", "• ")) for l in non_empty)):
            for line in lines:
                line = line.strip()
                if not line:
                    continue
                if line.startswith("###"):
                    add_sub_subsection_heading(doc, clean_markdown(line.lstrip("#").strip()))
                elif line.startswith("##"):
                    add_subsection_heading(doc, clean_markdown(line.lstrip("#").strip()))
                elif line.startswith("#"):
                    continue
                elif line.startswith(("- ", "* ", "• ")):
                    cleaned = _re.sub(r'^\s*[-•*]\s+', '', line).strip()
                    if cleaned:
                        _render_bullet_paragraph(doc, cleaned)
                else:
                    _render_formatted_paragraph(doc, line)
            continue

        # Regular paragraph
        _render_formatted_paragraph(doc, para_text)


def _strip_full_bold(text: str) -> str:
    """Strip excessive bold markers that make the whole line bold.

    Handles two LLM patterns:
      1. Single wrapper:   **entire line is bold**
      2. Multi-wrapper:    **Sentence one.** **Sentence two.** **Sentence three.**
         (every phrase individually wrapped → entire text looks bold)

    If >80% of visible text is inside ** markers, we strip ALL markers and
    return plain text so only genuine sub-phrase emphasis is rendered bold.
    """
    import re as _re
    stripped = text.strip()
    if "**" not in stripped:
        return text

    # Case 1: single wrapper — **entire text**
    if stripped.startswith("**") and stripped.endswith("**"):
        inner = stripped[2:-2]
        if "**" not in inner:
            return inner

    # Case 2: multiple bold chunks covering most of the text
    bold_chunks = _re.findall(r'\*\*(.*?)\*\*', stripped)
    if bold_chunks:
        bold_len = sum(len(c) for c in bold_chunks)
        plain = _re.sub(r'\*\*.*?\*\*', '', stripped)
        plain_len = len(plain.strip())
        total_visible = bold_len + plain_len
        if total_visible > 0 and bold_len / total_visible > 0.75:
            # Most of the text is bold — strip all markers
            return _re.sub(r'\*\*(.*?)\*\*', r'\1', stripped)

    return text


def _add_run_or_citation(para, text: str, bold: bool = False, italic: bool = False):
    """Add text as runs, stripping any [src_xxx_nnn] inline citations.

    Citations are not shown inline — all sources are listed in the bibliography.
    """
    import re

    if not text:
        return

    # Strip any inline citation references (LLM may still produce them occasionally)
    cleaned = re.sub(r'\s*\[src_[a-z0-9_]+\]', '', text)
    if not cleaned:
        return

    run = para.add_run(cleaned)
    if bold:
        run.bold = True
    if italic:
        run.italic = True


def _render_formatted_paragraph(doc: Document, text: str):
    """Add a body paragraph with **bold** and *italic* rendered as docx runs."""
    import re

    text = _strip_full_bold(text)
    para = doc.add_paragraph(style=STYLE_BODY)
    parts = re.split(r'(\*\*.*?\*\*)', text)
    for part in parts:
        if part.startswith("**") and part.endswith("**"):
            _add_run_or_citation(para, part[2:-2], bold=True)
        else:
            sub_parts = re.split(r'(\*[^*]+?\*)', part)
            for sub in sub_parts:
                if sub.startswith("*") and sub.endswith("*") and not sub.startswith("**"):
                    _add_run_or_citation(para, sub[1:-1], italic=True)
                else:
                    _add_run_or_citation(para, sub)


def _render_bullet_paragraph(doc: Document, text: str):
    """Add a presentation-style bulleted paragraph with colored accent marker."""
    import re

    text = _strip_full_bold(text)
    para = doc.add_paragraph(style=STYLE_BODY)
    para.paragraph_format.left_indent = Cm(0.9)
    para.paragraph_format.first_line_indent = Cm(-0.5)
    para.paragraph_format.space_before = Pt(4)
    para.paragraph_format.space_after = Pt(4)
    para.paragraph_format.line_spacing = Pt(17)

    # Colored accent bullet
    bullet_run = para.add_run("\u25CF  ")
    bullet_run.font.size = Pt(10)
    bullet_run.font.color.rgb = GOLD

    parts = re.split(r'(\*\*.*?\*\*)', text)
    for part in parts:
        if part.startswith("**") and part.endswith("**"):
            r = para.add_run(part[2:-2])
            r.bold = True
            r.font.size = Pt(10.5)
            r.font.name = FONT_BODY
            r.font.color.rgb = DARK_TEXT
        else:
            sub_parts = re.split(r'(\*[^*]+?\*)', part)
            for sub in sub_parts:
                if sub.startswith("*") and sub.endswith("*") and not sub.startswith("**"):
                    r = para.add_run(sub[1:-1])
                    r.italic = True
                    r.font.size = Pt(10.5)
                    r.font.name = FONT_BODY
                    r.font.color.rgb = STEEL_GRAY
                else:
                    _add_run_or_citation(para, sub)


def add_markdown_table(doc: Document, table_md: str):
    """Render a markdown table in docx with styled headers and bold support."""
    import re

    lines = [l.strip() for l in table_md.strip().split("\n") if l.strip()]
    # Keep only pipe-delimited lines; drop separator rows and stray non-table text
    pipe_lines = [l for l in lines if "|" in l]
    data_lines = [l for l in pipe_lines if not all(c in "-| :" for c in l)]

    if len(data_lines) < 2:
        add_body_text(doc, clean_markdown(table_md))
        return

    rows = []
    for line in data_lines:
        cells = [c.strip() for c in line.split("|") if c.strip()]
        rows.append(cells)

    if not rows:
        return

    n_cols = max(len(r) for r in rows)
    n_rows = len(rows)

    table = doc.add_table(rows=n_rows, cols=n_cols)
    table.style = "Table Grid"
    _keep_table_on_one_page(table)

    for i, row_data in enumerate(rows):
        for j, cell_text in enumerate(row_data):
            if j >= n_cols:
                break
            cell = table.cell(i, j)
            cell.text = ""

            if i == 0:
                # Header row — navy bg, white bold text
                set_cell_shading(cell, "006B77")
                para = cell.paragraphs[0]
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = para.add_run(clean_markdown(cell_text))
                run.font.color.rgb = WHITE
                run.font.bold = True
                run.font.size = Pt(9)
                run.font.name = FONT_BODY
            else:
                # Data row — very subtle alternation
                if i % 2 == 0:
                    set_cell_shading(cell, "F0F7F7")
                para = cell.paragraphs[0]
                parts = re.split(r'(\*\*.*?\*\*)', cell_text)
                for part in parts:
                    if part.startswith("**") and part.endswith("**"):
                        run = para.add_run(part[2:-2])
                        run.bold = True
                        run.font.size = Pt(8)
                        run.font.name = FONT_BODY
                    else:
                        run = para.add_run(part)
                        run.font.size = Pt(8)
                        run.font.name = FONT_BODY

    doc.add_paragraph("")


# ─── PPTX-Style Layout Components ───────────────────────────────────────────


def add_kpi_card_row(doc: Document, metrics: list[dict]):
    """Add PPTX-style KPI dashboard: dark teal cards with white/bright numbers.

    Args:
        metrics: [{"label": "Market Size", "value": "US$ 45.2 Bn",
                   "subtitle": "by 2032"}]
    """
    n = len(metrics)
    if n == 0:
        return

    table = doc.add_table(rows=1, cols=n)
    _remove_table_borders(table)
    _keep_table_on_one_page(table)
    table.alignment = 1  # CENTER

    for j, metric in enumerate(metrics):
        cell = table.cell(0, j)
        cell.text = ""

        # Dark teal background
        set_cell_shading(cell, "006B77")
        _set_generous_padding(cell, top=180, bottom=180, left=160, right=160)

        # Small label: Calibri Light, WHITE, ALL-CAPS
        para = cell.paragraphs[0]
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        label_run = para.add_run(metric.get("label", "").upper())
        label_run.font.name = FONT_LABEL
        label_run.font.size = Pt(8)
        label_run.font.color.rgb = WHITE
        label_run.font.bold = True

        # Big value: Georgia, WHITE (bright, high contrast on dark bg)
        val_para = cell.add_paragraph()
        val_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        val_para.paragraph_format.space_before = Pt(6)
        val_para.paragraph_format.space_after = Pt(2)
        val_run = val_para.add_run(metric.get("value", "—"))
        val_run.font.name = FONT_DISPLAY
        val_run.font.size = Pt(24)
        val_run.font.bold = True
        val_run.font.color.rgb = WHITE

        # Subtitle: Calibri Light, light teal-green (visible on dark bg)
        subtitle = metric.get("subtitle", "")
        if subtitle:
            sub_para = cell.add_paragraph()
            sub_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            sub_para.paragraph_format.space_before = Pt(0)
            sub_para.paragraph_format.space_after = Pt(0)
            sub_run = sub_para.add_run(subtitle)
            sub_run.font.name = FONT_LABEL
            sub_run.font.size = Pt(8)
            sub_run.font.color.rgb = LIGHT_TEAL_GREEN
            sub_run.font.italic = True

    doc.add_paragraph("")


def add_kpi_dashboard(doc: Document, metrics: list[dict]):
    """Legacy KPI dashboard — redirects to new PPTX-style card row."""
    add_kpi_card_row(doc, metrics)


def add_accent_bar_card(doc: Document, title: str, body: str,
                        accent_color: str = "00BCD4"):
    """Card with thin colored left border — PPTX accent-bar pattern."""
    table = doc.add_table(rows=1, cols=1)
    _remove_table_borders(table)
    _keep_table_on_one_page(table)
    table.alignment = 1  # CENTER
    cell = table.cell(0, 0)
    cell.text = ""
    set_cell_shading(cell, "FFFFFF")
    _add_left_accent_border(cell, accent_color, thickness="36")
    _set_generous_padding(cell, top=160, bottom=160, left=200, right=200)

    # Title in navy serif
    para = cell.paragraphs[0]
    run = para.add_run(title)
    run.font.name = FONT_DISPLAY
    run.font.size = Pt(12)
    run.font.bold = True
    run.font.color.rgb = NAVY

    # Body text
    if body:
        body_para = cell.add_paragraph()
        body_para.paragraph_format.space_before = Pt(4)
        run = body_para.add_run(body)
        run.font.name = FONT_BODY
        run.font.size = Pt(10)
        run.font.color.rgb = DARK_TEXT

    doc.add_paragraph("")


def add_two_column_panel(doc: Document, left_title: str, left_body: str,
                         right_title: str, right_body: str,
                         left_bg: str = "006B77", right_bg: str = "E8F4F5"):
    """Two-column panel layout for side-by-side content."""
    table = doc.add_table(rows=1, cols=2)
    _remove_table_borders(table)
    _keep_table_on_one_page(table)
    table.alignment = 1

    panels = [(left_title, left_body, left_bg), (right_title, right_body, right_bg)]
    for col_idx, (title, body, bg) in enumerate(panels):
        cell = table.cell(0, col_idx)
        cell.text = ""
        set_cell_shading(cell, bg)
        _set_generous_padding(cell, top=180, bottom=180, left=200, right=200)

        is_dark = bg in ("006B77", "003F48", "5A7D8C")
        title_color = GOLD if is_dark else NAVY
        text_color = WHITE if is_dark else DARK_TEXT

        # Title
        para = cell.paragraphs[0]
        run = para.add_run(title)
        run.font.name = FONT_DISPLAY
        run.font.size = Pt(13)
        run.font.bold = True
        run.font.color.rgb = title_color

        # Body
        if body:
            body_para = cell.add_paragraph()
            body_para.paragraph_format.space_before = Pt(6)
            run = body_para.add_run(body)
            run.font.name = FONT_BODY
            run.font.size = Pt(10)
            run.font.color.rgb = text_color

    doc.add_paragraph("")


def add_footer_insight_bar(doc: Document, text: str, bg_color: str = "006B77"):
    """Full-width colored bar with key takeaway — PPTX footer pattern."""
    table = doc.add_table(rows=1, cols=1)
    _remove_table_borders(table)
    _keep_table_on_one_page(table)
    table.alignment = 1
    cell = table.cell(0, 0)
    cell.text = ""
    set_cell_shading(cell, bg_color)
    _set_generous_padding(cell, top=140, bottom=140, left=200, right=200)

    para = cell.paragraphs[0]

    # "KEY TAKEAWAY" label — bright on dark bg, teal on light bg
    is_dark = bg_color in ("006B77", "003F48", "009688", "2E7D32")
    label = para.add_run("KEY TAKEAWAY  ")
    label.font.name = FONT_LABEL
    label.font.size = Pt(8)
    label.font.bold = True
    label.font.color.rgb = LIGHT_TEAL_GREEN if is_dark else GOLD

    # Takeaway text — white on dark bg, dark text on light bg
    body_color = WHITE if is_dark else DARK_TEXT
    body = para.add_run(text)
    body.font.name = FONT_BODY
    body.font.size = Pt(10)
    body.font.color.rgb = body_color

    doc.add_paragraph("")


# ─── Spatial Layout Helpers (PPTX-grade 2-column panels) ────────────────────


def add_chart_with_sidebar(doc: Document, chart_img_stream, build_sidebar_fn,
                           chart_width: float = 5.5, sidebar_width: float = 4.0,
                           caption: str = ""):
    """2-column panel: chart image on left, sidebar cards on right.

    Replicates the PPTX 60/40 split pattern where a bar/line chart occupies
    the left panel and KPI/driver cards fill the right panel.

    Args:
        chart_img_stream: BytesIO returned by any chart_* function
        build_sidebar_fn: callable(cell) that populates right sidebar cell
        chart_width: left column width in inches (default 5.5)
        sidebar_width: right column width in inches (default 4.0)
        caption: optional chart caption rendered below the panel
    """
    table = doc.add_table(rows=1, cols=2)
    _remove_table_borders(table)
    _keep_table_on_one_page(table)

    # Fix column widths (landscape usable width ≈ 9.5")
    table.columns[0].width = Inches(chart_width)
    table.columns[1].width = Inches(sidebar_width)

    # ── Left: chart image ────────────────────────────────────────────────────
    left_cell = table.cell(0, 0)
    left_cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
    _set_generous_padding(left_cell, top=60, bottom=60, left=0, right=100)

    left_para = left_cell.paragraphs[0]
    left_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = left_para.add_run()

    if hasattr(chart_img_stream, "seek"):
        chart_img_stream.seek(0)
    run.add_picture(chart_img_stream, width=Inches(chart_width - 0.18))

    # ── Right: caller-provided sidebar content ───────────────────────────────
    right_cell = table.cell(0, 1)
    right_cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
    build_sidebar_fn(right_cell)

    if caption:
        doc.add_paragraph(caption, style=STYLE_CHART_CAPTION)
    else:
        doc.add_paragraph("")


def add_cagr_analysis_card(cell, cagr_pct, from_val: str = "", to_val: str = "",
                           label: str = "CAGR"):
    """Populate a sidebar cell with a CAGR analysis card (navy bg, gold value).

    Replicates PPTX slide 4 top-right card: label, large CAGR %, from→to range.

    Args:
        cell: table cell to populate
        cagr_pct: CAGR string e.g. "8.4%" or float 0.084
        from_val: start year/value label e.g. "US$ 12.5 Bn (2023)"
        to_val: end year/value label e.g. "US$ 28.1 Bn (2032)"
        label: card label (default "CAGR")
    """
    set_cell_shading(cell, "006B77")
    _set_generous_padding(cell, top=140, bottom=140, left=140, right=100)

    # Format cagr_pct value
    if isinstance(cagr_pct, float):
        pct_str = f"{cagr_pct * 100:.1f}%"
    else:
        pct_str = str(cagr_pct) if cagr_pct else "—"

    # Label — Calibri Light, WHITE, ALL-CAPS
    first_para = cell.paragraphs[0]
    first_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    label_run = first_para.add_run(label.upper())
    label_run.font.name = FONT_LABEL
    label_run.font.size = Pt(8)
    label_run.font.bold = True
    label_run.font.color.rgb = WHITE

    # Large CAGR value — Georgia, WHITE, 28pt (high contrast on dark bg)
    val_para = cell.add_paragraph()
    val_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    val_para.paragraph_format.space_before = Pt(6)
    val_para.paragraph_format.space_after = Pt(6)
    val_run = val_para.add_run(pct_str)
    val_run.font.name = FONT_DISPLAY
    val_run.font.size = Pt(28)
    val_run.font.bold = True
    val_run.font.color.rgb = WHITE

    # Progress bar — gold fill / medium-navy remainder (PPTX slide 4 style)
    try:
        cagr_f = float(pct_str.replace("%", "").strip())
        max_scale = max(15.0, cagr_f * 2.0)
        fill_ratio = min(cagr_f / max_scale, 1.0)

        bar_total = 18
        filled = max(1, round(fill_ratio * bar_total))
        empty = bar_total - filled

        bar_para = cell.add_paragraph()
        bar_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        bar_para.paragraph_format.space_before = Pt(8)
        bar_para.paragraph_format.space_after = Pt(2)

        filled_run = bar_para.add_run("\u2588" * filled)
        filled_run.font.color.rgb = LIGHT_TEAL_GREEN
        filled_run.font.size = Pt(9)
        filled_run.font.name = FONT_BODY

        empty_run = bar_para.add_run("\u2591" * empty)
        empty_run.font.color.rgb = RGBColor(0x00, 0x8B, 0x96)
        empty_run.font.size = Pt(9)
        empty_run.font.name = FONT_BODY
    except Exception:
        pass

    # From → To range — Calibri Light, LIGHT_TEAL_GREEN (visible on dark bg)
    if from_val or to_val:
        range_para = cell.add_paragraph()
        range_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        range_para.paragraph_format.space_before = Pt(4)
        range_run = range_para.add_run(f"{from_val}  \u2192  {to_val}")
        range_run.font.name = FONT_LABEL
        range_run.font.size = Pt(8)
        range_run.font.color.rgb = LIGHT_TEAL_GREEN


def add_numbered_drivers_card(cell, items: list, title: str = ""):
    """Populate a sidebar cell with a numbered driver/insight list.

    Replicates PPTX slide 4 right-bottom panel: circled numbers + title + desc.

    Args:
        cell: table cell to populate
        items: list of {"title": str, "desc": str} or plain strings
        title: optional section label (e.g. "KEY GROWTH DRIVERS")
    """
    _set_generous_padding(cell, top=100, bottom=100, left=120, right=80)

    CIRCLE_BADGES = ["\u2460", "\u2461", "\u2462", "\u2463", "\u2464", "\u2465"]

    first_para = cell.paragraphs[0]
    # If cell already has content (e.g. a CAGR card above), always start a new paragraph
    cell_already_has_content = bool(first_para.runs) or len(cell.paragraphs) > 1
    if title:
        title_para = cell.add_paragraph() if cell_already_has_content else first_para
        title_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        if cell_already_has_content:
            title_para.paragraph_format.space_before = Pt(10)
        title_run = title_para.add_run(title.upper())
        title_run.font.name = FONT_LABEL
        title_run.font.size = Pt(8)
        title_run.font.bold = True
        title_run.font.color.rgb = LIGHT_TEAL_GREEN
        rPr = title_run._element.get_or_add_rPr()
        spacing = OxmlElement("w:spacing")
        spacing.set(qn("w:val"), "40")
        rPr.append(spacing)
    else:
        if not cell_already_has_content:
            first_para.clear()

    for i, item in enumerate(items[:6]):
        badge = CIRCLE_BADGES[i] if i < len(CIRCLE_BADGES) else f"{i + 1}."
        if isinstance(item, dict):
            item_title = item.get("title", "")
            item_desc = item.get("desc", "")
        else:
            item_title = str(item)
            item_desc = ""

        # Badge + title on one line (white text for dark sidebar bg)
        item_para = cell.add_paragraph()
        item_para.paragraph_format.space_before = Pt(9 if i == 0 else 6)
        item_para.paragraph_format.space_after = Pt(0)

        badge_run = item_para.add_run(f"{badge}  ")
        badge_run.font.name = FONT_BODY
        badge_run.font.size = Pt(12)
        badge_run.font.bold = True
        badge_run.font.color.rgb = LIGHT_TEAL_GREEN

        title_run = item_para.add_run(item_title)
        title_run.font.name = FONT_BODY
        title_run.font.size = Pt(10)
        title_run.font.bold = True
        title_run.font.color.rgb = WHITE

        if item_desc:
            desc_para = cell.add_paragraph()
            desc_para.paragraph_format.space_before = Pt(1)
            desc_para.paragraph_format.space_after = Pt(0)
            desc_para.paragraph_format.left_indent = Inches(0.28)
            desc_run = desc_para.add_run(item_desc)
            desc_run.font.name = FONT_BODY
            desc_run.font.size = Pt(9)
            desc_run.font.color.rgb = LIGHT_TEAL_GREEN


def add_kpi_cards_with_badge(doc: Document, metrics: list):
    """KPI card row with optional CAGR pill badges (enhanced add_kpi_card_row).

    Args:
        metrics: list of {"label": str, "value": str, "subtitle": str,
                          "badge": str, "badge_color": "#RRGGBB"} dicts
                 "badge" and "badge_color" are optional.
    """
    n = len(metrics)
    if n == 0:
        return

    table = doc.add_table(rows=1, cols=n)
    _remove_table_borders(table)
    _keep_table_on_one_page(table)
    table.alignment = 1  # CENTER

    for j, metric in enumerate(metrics):
        cell = table.cell(0, j)
        cell.text = ""
        set_cell_shading(cell, "006B77")
        _set_generous_padding(cell, top=160, bottom=160, left=140, right=140)

        # Label — Calibri Light, WHITE, ALL-CAPS
        para = cell.paragraphs[0]
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        label_run = para.add_run(metric.get("label", "").upper())
        label_run.font.name = FONT_LABEL
        label_run.font.size = Pt(7)
        label_run.font.color.rgb = WHITE
        label_run.font.bold = True

        # Value — Georgia, WHITE, 22pt (high contrast on dark teal)
        val_para = cell.add_paragraph()
        val_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        val_para.paragraph_format.space_before = Pt(5)
        val_para.paragraph_format.space_after = Pt(2)
        val_run = val_para.add_run(metric.get("value", "—"))
        val_run.font.name = FONT_DISPLAY
        val_run.font.size = Pt(22)
        val_run.font.bold = True
        val_run.font.color.rgb = WHITE

        # Subtitle — Calibri Light, LIGHT_TEAL_GREEN (visible on dark bg)
        subtitle = metric.get("subtitle", "")
        if subtitle:
            sub_para = cell.add_paragraph()
            sub_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            sub_para.paragraph_format.space_before = Pt(0)
            sub_para.paragraph_format.space_after = Pt(4)
            sub_run = sub_para.add_run(subtitle)
            sub_run.font.name = FONT_LABEL
            sub_run.font.size = Pt(8)
            sub_run.font.color.rgb = LIGHT_TEAL_GREEN

        # Badge pill — e.g. "▲ 5.6% p.a." with thin separator line above
        badge_text = metric.get("badge", "")
        if badge_text:
            badge_color_hex = metric.get("badge_color", "00BCD4").lstrip("#")
            try:
                badge_rgb = RGBColor.from_string(badge_color_hex)
            except Exception:
                badge_rgb = GOLD

            # Thin separator line (gold dots)
            sep_para = cell.add_paragraph()
            sep_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            sep_para.paragraph_format.space_before = Pt(4)
            sep_para.paragraph_format.space_after = Pt(0)
            sep_run = sep_para.add_run("\u2500" * 18)
            sep_run.font.size = Pt(4)
            sep_run.font.name = FONT_BODY
            sep_run.font.color.rgb = badge_rgb

            badge_para = cell.add_paragraph()
            badge_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            badge_para.paragraph_format.space_before = Pt(2)
            badge_run = badge_para.add_run(f"  {badge_text}  ")
            badge_run.font.name = FONT_LABEL
            badge_run.font.size = Pt(7)
            badge_run.font.bold = True
            try:
                badge_run.font.color.rgb = RGBColor.from_string(badge_color_hex)
            except Exception:
                badge_run.font.color.rgb = GOLD

    doc.add_paragraph("")


# ─── Visual Elements ─────────────────────────────────────────────────────────


def add_callout_box(doc: Document, text: str, box_type: str = "insight",
                    title: str = None):
    """Add a consulting-grade callout box with colored left accent bar."""
    BOX_STYLES = {
        "insight":     {"bar": "00BCD4", "bg": "E8F4F5", "icon": "\u2139",  "default_title": "Key Insight"},
        "warning":     {"bar": "5A7D8C", "bg": "E8F4F5", "icon": "\u25A0",  "default_title": "Note"},
        "success":     {"bar": "00BCD4", "bg": "E8F4F5", "icon": "\u25B2",  "default_title": "Highlight"},
        "key_finding": {"bar": "006B77", "bg": "E8F4F5", "icon": "\u2605",  "default_title": "Key Finding"},
    }
    style = BOX_STYLES.get(box_type, BOX_STYLES["insight"])
    display_title = title or style["default_title"]

    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    _keep_table_on_one_page(table)
    table.alignment = 1

    cell = table.cell(0, 0)
    cell.text = ""
    set_cell_shading(cell, style["bg"])

    # Generous padding
    _set_generous_padding(cell, top=200, bottom=200, left=240, right=240)

    # Thick left border (accent bar)
    tcPr = cell._element.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), "48")
    left.set(qn("w:space"), "0")
    left.set(qn("w:color"), style["bar"])
    tcBorders.append(left)
    for edge in ("top", "right", "bottom"):
        e = OxmlElement(f"w:{edge}")
        e.set(qn("w:val"), "single")
        e.set(qn("w:sz"), "4")
        e.set(qn("w:space"), "0")
        e.set(qn("w:color"), "E8E8E8")
        tcBorders.append(e)
    tcPr.append(tcBorders)

    # Title line: icon + title
    para = cell.paragraphs[0]
    icon_run = para.add_run(f"{style['icon']}  ")
    icon_run.font.size = Pt(14)
    icon_run.font.color.rgb = RGBColor.from_string(style["bar"])
    icon_run.font.name = FONT_DISPLAY

    title_run = para.add_run(display_title)
    title_run.font.size = Pt(11)
    title_run.font.bold = True
    title_run.font.color.rgb = RGBColor.from_string(style["bar"])
    title_run.font.name = FONT_DISPLAY

    # Body text
    body_para = cell.add_paragraph()
    body_para.paragraph_format.space_before = Pt(4)
    body_para.paragraph_format.space_after = Pt(0)

    import re
    text = _strip_full_bold(text)
    parts = re.split(r'(\*\*.*?\*\*)', text)
    for part in parts:
        if part.startswith("**") and part.endswith("**"):
            r = body_para.add_run(part[2:-2])
            r.bold = True
            r.font.size = Pt(10)
            r.font.name = FONT_BODY
            r.font.color.rgb = DARK_TEXT
        else:
            r = body_para.add_run(part)
            r.font.size = Pt(10)
            r.font.name = FONT_BODY
            r.font.color.rgb = DARK_TEXT

    doc.add_paragraph("")


def add_chart_image(doc: Document, image_stream, width=Inches(9.0), caption: str = ""):
    """Insert a chart image from BytesIO — full-width for slide-style layout."""
    doc.add_picture(image_stream, width=width)
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    if caption:
        doc.add_paragraph(caption, style=STYLE_CHART_CAPTION)


def add_insight_title(doc: Document, insight: str, subtitle: str = ""):
    """Add a large insight-driven headline — Georgia serif, navy."""
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(16)
    para.paragraph_format.space_after = Pt(4)
    run = para.add_run(insight)
    run.font.size = Pt(18)
    run.font.bold = True
    run.font.color.rgb = NAVY
    run.font.name = FONT_DISPLAY

    if subtitle:
        sub_para = doc.add_paragraph()
        sub_para.paragraph_format.space_before = Pt(0)
        sub_para.paragraph_format.space_after = Pt(10)
        sub_run = sub_para.add_run(subtitle)
        sub_run.font.size = Pt(10)
        sub_run.font.color.rgb = STEEL_GRAY
        sub_run.font.name = FONT_LABEL
        sub_run.font.italic = True


def add_slide_divider(doc: Document, color_hex: str = "00BCD4"):
    """Add a gold divider line as a visual section break."""
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(20)
    para.paragraph_format.space_after = Pt(16)

    pPr = para._element.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "24")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), color_hex)
    pBdr.append(bottom)
    pPr.append(pBdr)


def set_cell_shading(cell, color_hex: str):
    """Set background shading on a table cell."""
    shading = cell._element.get_or_add_tcPr()
    shading_elm = shading.makeelement(qn("w:shd"), {
        qn("w:fill"): color_hex,
        qn("w:val"): "clear",
    })
    shading.append(shading_elm)


def add_hyperlink(paragraph, text: str, url: str, color: RGBColor = None,
                  font_size: Pt = None):
    """Add a clickable hyperlink to an existing paragraph."""
    if color is None:
        color = GOLD

    part = paragraph.part
    r_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )

    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)

    run_element = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")

    c_element = OxmlElement("w:color")
    c_element.set(qn("w:val"), str(color))
    rPr.append(c_element)

    u_element = OxmlElement("w:u")
    u_element.set(qn("w:val"), "single")
    rPr.append(u_element)

    rFonts = OxmlElement("w:rFonts")
    rFonts.set(qn("w:ascii"), FONT_BODY)
    rFonts.set(qn("w:hAnsi"), FONT_BODY)
    rPr.append(rFonts)

    if font_size:
        sz = OxmlElement("w:sz")
        sz.set(qn("w:val"), str(int(font_size.pt * 2)))
        rPr.append(sz)

    run_element.append(rPr)

    t_element = OxmlElement("w:t")
    t_element.set(qn("xml:space"), "preserve")
    t_element.text = text
    run_element.append(t_element)

    hyperlink.append(run_element)
    paragraph._element.append(hyperlink)

    return hyperlink


def add_horizontal_rule(doc: Document, color: RGBColor = None, thickness: int = 12):
    """Add a colored horizontal rule using a bottom border."""
    if color is None:
        color = NAVY

    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(0)
    para.paragraph_format.space_after = Pt(6)

    pPr = para._element.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), str(thickness))
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), str(color))
    pBdr.append(bottom)
    pPr.append(pBdr)

    return para


# ─── Presentation-Grade Visual Components ───────────────────────────────────


def add_icon_feature_cards(doc: Document, features: list, cols: int = 3):
    """Add a row of icon + title + description feature cards.

    Each feature is a card with a colored icon, bold title, and short description.
    Great for showing drivers, opportunities, or key points visually.

    Args:
        features: list of {"icon": "◉", "title": str, "desc": str, "color": hex}
        cols: number of columns (2-4)
    """
    if not features:
        return

    n = min(len(features), cols)
    table = doc.add_table(rows=1, cols=n)
    _remove_table_borders(table)
    _keep_table_on_one_page(table)
    table.alignment = 1

    card_bgs = ["006B77", "009688", "00BCD4", "2E7D32"]

    for j, feat in enumerate(features[:n]):
        cell = table.cell(0, j)
        cell.text = ""
        bg = feat.get("bg", card_bgs[j % len(card_bgs)])
        set_cell_shading(cell, bg)
        _set_generous_padding(cell, top=200, bottom=200, left=180, right=180)

        # Icon (large)
        icon_para = cell.paragraphs[0]
        icon_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        icon_run = icon_para.add_run(feat.get("icon", "◉"))
        icon_run.font.name = FONT_BODY
        icon_run.font.size = Pt(28)
        icon_run.font.color.rgb = WHITE

        # Title
        title_para = cell.add_paragraph()
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_para.paragraph_format.space_before = Pt(8)
        title_para.paragraph_format.space_after = Pt(4)
        title_run = title_para.add_run(feat.get("title", ""))
        title_run.font.name = FONT_DISPLAY
        title_run.font.size = Pt(12)
        title_run.font.bold = True
        title_run.font.color.rgb = WHITE

        # Description
        desc = feat.get("desc", "")
        if desc:
            desc_para = cell.add_paragraph()
            desc_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            desc_para.paragraph_format.space_before = Pt(0)
            desc_run = desc_para.add_run(desc)
            desc_run.font.name = FONT_BODY
            desc_run.font.size = Pt(9)
            desc_run.font.color.rgb = OFF_WHITE

    # Handle overflow — additional rows
    if len(features) > n:
        doc.add_paragraph("")
        remaining = features[n:]
        add_icon_feature_cards(doc, remaining, cols)
    else:
        doc.add_paragraph("")


def add_stat_highlight_row(doc: Document, stats: list):
    """Add a row of large stat numbers with labels — great for key metrics.

    Lighter background than KPI cards, more presentation-friendly.

    Args:
        stats: list of {"number": "45.2 Bn", "label": "Market Size", "sublabel": "by 2032"}
    """
    if not stats:
        return

    n = min(len(stats), 4)
    table = doc.add_table(rows=1, cols=n)
    _remove_table_borders(table)
    _keep_table_on_one_page(table)
    table.alignment = 1

    for j, stat in enumerate(stats[:n]):
        cell = table.cell(0, j)
        cell.text = ""
        set_cell_shading(cell, "E8F4F5")
        _set_generous_padding(cell, top=200, bottom=200, left=120, right=120)

        # Large number
        num_para = cell.paragraphs[0]
        num_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        num_run = num_para.add_run(stat.get("number", "—"))
        num_run.font.name = FONT_DISPLAY
        num_run.font.size = Pt(32)
        num_run.font.bold = True
        num_run.font.color.rgb = NAVY

        # Label
        label_para = cell.add_paragraph()
        label_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        label_para.paragraph_format.space_before = Pt(4)
        label_para.paragraph_format.space_after = Pt(2)
        label_run = label_para.add_run(stat.get("label", "").upper())
        label_run.font.name = FONT_LABEL
        label_run.font.size = Pt(9)
        label_run.font.bold = True
        label_run.font.color.rgb = STEEL_GRAY

        # Sublabel
        sublabel = stat.get("sublabel", "")
        if sublabel:
            sub_para = cell.add_paragraph()
            sub_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            sub_para.paragraph_format.space_before = Pt(0)
            sub_run = sub_para.add_run(sublabel)
            sub_run.font.name = FONT_LABEL
            sub_run.font.size = Pt(8)
            sub_run.font.color.rgb = STEEL_GRAY

    doc.add_paragraph("")


def add_two_column_content(doc: Document, left_content: str, right_content: str,
                           left_title: str = "", right_title: str = ""):
    """Two-column layout for side-by-side bullet content.

    Both columns have a light background. Renders markdown bullet points.
    """
    table = doc.add_table(rows=1, cols=2)
    _remove_table_borders(table)
    _keep_table_on_one_page(table)
    table.alignment = 1
    table.columns[0].width = Inches(4.6)
    table.columns[1].width = Inches(4.6)

    for col_idx, (title, content) in enumerate(
        [(left_title, left_content), (right_title, right_content)]
    ):
        cell = table.cell(0, col_idx)
        cell.text = ""
        bg = "E8F4F5" if col_idx == 0 else "F0F7F7"
        set_cell_shading(cell, bg)
        _add_left_accent_border(cell, "00BCD4" if col_idx == 0 else "009688", thickness="36")
        _set_generous_padding(cell, top=180, bottom=180, left=200, right=160)

        # Title
        if title:
            para = cell.paragraphs[0]
            run = para.add_run(title.upper())
            run.font.name = FONT_LABEL
            run.font.size = Pt(9)
            run.font.bold = True
            run.font.color.rgb = NAVY
            # Add letter spacing
            rPr = run._element.get_or_add_rPr()
            spacing = OxmlElement("w:spacing")
            spacing.set(qn("w:val"), "30")
            rPr.append(spacing)

        # Render content as bullets in-cell
        if content:
            import re as _re
            lines = content.strip().split("\n")
            for line in lines:
                line = _re.sub(r'^\s*[-•*]\s+', '', line).strip()
                if not line:
                    continue
                bp = cell.add_paragraph()
                bp.paragraph_format.space_before = Pt(3)
                bp.paragraph_format.space_after = Pt(2)
                bp.paragraph_format.left_indent = Cm(0.4)
                bp.paragraph_format.first_line_indent = Cm(-0.3)

                bullet_run = bp.add_run("\u2022  ")
                bullet_run.font.size = Pt(10)
                bullet_run.font.color.rgb = GOLD

                # Handle bold
                line = _strip_full_bold(line)
                parts = _re.split(r'(\*\*.*?\*\*)', line)
                for part in parts:
                    if part.startswith("**") and part.endswith("**"):
                        r = bp.add_run(part[2:-2])
                        r.bold = True
                        r.font.size = Pt(9.5)
                        r.font.name = FONT_BODY
                        r.font.color.rgb = DARK_TEXT
                    else:
                        r = bp.add_run(part)
                        r.font.size = Pt(9.5)
                        r.font.name = FONT_BODY
                        r.font.color.rgb = DARK_TEXT

    doc.add_paragraph("")


def add_company_profile_card(doc: Document, company_name: str, profile_text: str):
    """Add a visual company profile card with accent bar and structured content.

    Renders as a card with navy left accent, company name header, and bullet content.
    """
    table = doc.add_table(rows=1, cols=1)
    _remove_table_borders(table)
    _keep_table_on_one_page(table)
    table.alignment = 1
    cell = table.cell(0, 0)
    cell.text = ""
    set_cell_shading(cell, "F0F7F7")
    _add_left_accent_border(cell, "006B77", thickness="48")
    _set_generous_padding(cell, top=180, bottom=180, left=220, right=200)

    # Company name header
    para = cell.paragraphs[0]
    run = para.add_run(company_name)
    run.font.name = FONT_DISPLAY
    run.font.size = Pt(14)
    run.font.bold = True
    run.font.color.rgb = NAVY

    # Accent line
    line_para = cell.add_paragraph()
    line_para.paragraph_format.space_before = Pt(4)
    line_para.paragraph_format.space_after = Pt(6)
    line_run = line_para.add_run("\u2500" * 30)
    line_run.font.size = Pt(4)
    line_run.font.color.rgb = GOLD

    # Profile content as bullets
    if profile_text:
        import re as _re
        lines = profile_text.strip().split("\n")
        for line in lines:
            line = _re.sub(r'^\s*[-•*]\s+', '', line).strip()
            if not line:
                continue
            bp = cell.add_paragraph()
            bp.paragraph_format.space_before = Pt(2)
            bp.paragraph_format.space_after = Pt(2)
            bp.paragraph_format.left_indent = Cm(0.4)
            bp.paragraph_format.first_line_indent = Cm(-0.3)

            bullet_run = bp.add_run("\u25B8  ")
            bullet_run.font.size = Pt(9)
            bullet_run.font.color.rgb = GOLD

            line = _strip_full_bold(line)
            parts = _re.split(r'(\*\*.*?\*\*)', line)
            for part in parts:
                if part.startswith("**") and part.endswith("**"):
                    r = bp.add_run(part[2:-2])
                    r.bold = True
                    r.font.size = Pt(9.5)
                    r.font.name = FONT_BODY
                    r.font.color.rgb = DARK_TEXT
                else:
                    r = bp.add_run(part)
                    r.font.size = Pt(9.5)
                    r.font.name = FONT_BODY
                    r.font.color.rgb = DARK_TEXT

    doc.add_paragraph("")


def add_visual_bullet_list(doc: Document, items: list, title: str = "",
                           accent_color: str = "00BCD4"):
    """Add a presentation-style bullet list with colored accent dots and larger text.

    More visual than regular paragraphs — each bullet has an accent-colored marker
    and slightly larger, bolder text suitable for slide display.

    Args:
        items: list of strings (can contain **bold** markdown)
        title: optional section title above the list
        accent_color: hex color for bullet markers
    """
    import re as _re

    if title:
        title_para = doc.add_paragraph()
        title_para.paragraph_format.space_before = Pt(12)
        title_para.paragraph_format.space_after = Pt(8)
        run = title_para.add_run(title)
        run.font.name = FONT_DISPLAY
        run.font.size = Pt(14)
        run.font.bold = True
        run.font.color.rgb = NAVY

    try:
        accent_rgb = RGBColor.from_string(accent_color)
    except Exception:
        accent_rgb = GOLD

    for item in items:
        if not item or not item.strip():
            continue
        # Clean bullet markers
        item = _re.sub(r'^\s*[-•*]\s+', '', item).strip()
        if not item:
            continue

        para = doc.add_paragraph()
        para.paragraph_format.space_before = Pt(4)
        para.paragraph_format.space_after = Pt(4)
        para.paragraph_format.left_indent = Cm(0.8)
        para.paragraph_format.first_line_indent = Cm(-0.5)
        para.paragraph_format.line_spacing = Pt(18)

        # Accent bullet marker
        bullet_run = para.add_run("\u25CF  ")
        bullet_run.font.size = Pt(11)
        bullet_run.font.color.rgb = accent_rgb

        # Render text with bold support
        item = _strip_full_bold(item)
        parts = _re.split(r'(\*\*.*?\*\*)', item)
        for part in parts:
            if part.startswith("**") and part.endswith("**"):
                r = para.add_run(part[2:-2])
                r.bold = True
                r.font.size = Pt(11)
                r.font.name = FONT_BODY
                r.font.color.rgb = DARK_TEXT
            else:
                r = para.add_run(part)
                r.font.size = Pt(11)
                r.font.name = FONT_BODY
                r.font.color.rgb = DARK_TEXT


def add_section_intro_strip(doc: Document, title: str, subtitle: str = "",
                            icon: str = "", metrics: list = None):
    """Add a visual section intro strip — large title with optional icon, subtitle, and metrics.

    Creates a presentation-like title card for a section. When *metrics* is provided,
    renders a 2-column layout: title on left, stacked metric badges on right.

    metrics: list of dicts with 'value' and 'label' keys (max 3).
    """
    if metrics:
        # ── Two-column layout: title | metrics ──────────────────────────
        table = doc.add_table(rows=1, cols=2)
        _remove_table_borders(table)
        _keep_table_on_one_page(table)
        table.alignment = 1
        _set_table_grid(table, [6400, 3400])

        # Left cell — title + subtitle
        left = table.cell(0, 0)
        left.text = ""
        set_cell_shading(left, "006B77")
        _set_generous_padding(left, top=240, bottom=240, left=300, right=100)

        para = left.paragraphs[0]
        if icon:
            icon_run = para.add_run(f"{icon}  ")
            icon_run.font.size = Pt(22)
            icon_run.font.color.rgb = GOLD

        title_run = para.add_run(title)
        title_run.font.name = FONT_DISPLAY
        title_run.font.size = Pt(20)
        title_run.font.bold = True
        title_run.font.color.rgb = WHITE

        if subtitle:
            sub_para = left.add_paragraph()
            sub_para.paragraph_format.space_before = Pt(6)
            sub_run = sub_para.add_run(subtitle)
            sub_run.font.name = FONT_LABEL
            sub_run.font.size = Pt(11)
            sub_run.font.color.rgb = LIGHT_TEAL_GREEN

        # Right cell — metric badges (darker panel)
        right = table.cell(0, 1)
        right.text = ""
        set_cell_shading(right, "004D56")
        _set_generous_padding(right, top=160, bottom=160, left=160, right=160)
        right.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

        for i, metric in enumerate(metrics[:3]):
            para = right.paragraphs[0] if i == 0 else right.add_paragraph()
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            if i > 0:
                para.paragraph_format.space_before = Pt(6)
                para.paragraph_format.space_after = Pt(0)

            val_run = para.add_run(str(metric.get("value", "")))
            val_run.font.name = FONT_DISPLAY
            val_run.font.size = Pt(15)
            val_run.font.bold = True
            val_run.font.color.rgb = GOLD

            label_run = para.add_run(f"  {metric.get('label', '')}")
            label_run.font.name = FONT_LABEL
            label_run.font.size = Pt(8)
            label_run.font.color.rgb = OFF_WHITE
    else:
        # ── Single-column full-width layout ─────────────────────────────
        table = doc.add_table(rows=1, cols=1)
        _remove_table_borders(table)
        _keep_table_on_one_page(table)
        table.alignment = 1
        cell = table.cell(0, 0)
        cell.text = ""
        set_cell_shading(cell, "006B77")
        _set_generous_padding(cell, top=240, bottom=240, left=300, right=200)

        para = cell.paragraphs[0]
        if icon:
            icon_run = para.add_run(f"{icon}  ")
            icon_run.font.size = Pt(22)
            icon_run.font.color.rgb = GOLD

        title_run = para.add_run(title)
        title_run.font.name = FONT_DISPLAY
        title_run.font.size = Pt(20)
        title_run.font.bold = True
        title_run.font.color.rgb = WHITE

        if subtitle:
            sub_para = cell.add_paragraph()
            sub_para.paragraph_format.space_before = Pt(6)
            sub_run = sub_para.add_run(subtitle)
            sub_run.font.name = FONT_LABEL
            sub_run.font.size = Pt(11)
            sub_run.font.color.rgb = LIGHT_TEAL_GREEN

    doc.add_paragraph("")


# ─── Bibliography ────────────────────────────────────────────────────────────


def add_bibliography_entry(doc: Document, citation_id: str, title: str,
                           publisher: str, date: str, url: str):
    """Add a formatted bibliography entry with a clickable URL."""
    para = doc.add_paragraph(style=STYLE_BIBLIOGRAPHY)

    run_id = para.add_run(f"[{citation_id}] ")
    run_id.bold = True
    run_id.font.size = Pt(9)
    run_id.font.name = FONT_BODY
    run_id.font.color.rgb = NAVY

    run_title = para.add_run(f"{title}. ")
    run_title.font.size = Pt(9)
    run_title.font.name = FONT_BODY

    date_str = date if date else "n.d."
    run_pub = para.add_run(f"{publisher}. {date_str}. ")
    run_pub.font.size = Pt(9)
    run_pub.font.name = FONT_BODY
    run_pub.font.color.rgb = STEEL_GRAY

    add_hyperlink(para, url, url, color=GOLD, font_size=Pt(8))

    return para
